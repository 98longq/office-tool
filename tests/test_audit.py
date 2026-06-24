from pathlib import Path
import sys
import tempfile
import unittest

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.shared import Cm

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from office_tool.audit import OfficialDocumentAuditor
from office_tool.config import OfficeToolConfig


class AuditTests(unittest.TestCase):
    def test_detects_meeting_minutes_structure(self):
        doc = Document()
        for text in [
            "内部资料不得外传",
            "会议纪要",
            "（12）",
            "某某单位办公室    2026年6月22日",
            "关于研究重点工作的会议纪要",
            "会议研究了有关事项。",
            "出席：张三、李四、王五",
            "分送：有关部门，有关单位。",
        ]:
            doc.add_paragraph(text)

        report = OfficialDocumentAuditor().audit_document(doc)
        roles = {element.name: element.text for element in report.elements}

        self.assertEqual(report.profile, "meeting_minutes")
        self.assertTrue(report.is_meeting_minutes)
        self.assertEqual(roles["meeting_number"], "（12）")
        self.assertEqual(roles["meeting_issue_line"], "某某单位办公室    2026年6月22日")
        self.assertEqual(roles["meeting_attendees"], "出席：张三、李四、王五")
        self.assertEqual(roles["distribution"], "分送：有关部门，有关单位。")
        codes = {finding.code for finding in report.findings}
        self.assertIn("meeting_red_rule_missing", codes)
        self.assertIn("distribution_layout_missing", codes)

    def test_detects_letter_head_from_document_number(self):
        doc = Document()
        for text in [
            "某某集团有限公司",
            "某部门函〔2026〕42号",
            "关于商请协助办理有关事项的函",
            "有关单位：",
            "请协助办理有关事项。",
            "某某集团有限公司",
            "2026年6月22日",
            "（内部资料　　不得外传）",
            "（联系人：张三；电话：12345678）",
        ]:
            doc.add_paragraph(text)

        report = OfficialDocumentAuditor().audit_document(doc)
        roles = {element.name: element.text for element in report.elements}

        self.assertEqual(report.profile, "letter_head")
        self.assertTrue(report.is_red_head)
        self.assertTrue(report.is_letter_head)
        self.assertEqual(roles["red_head"], "某某集团有限公司")
        self.assertEqual(roles["title"], "关于商请协助办理有关事项的函")
        self.assertEqual(roles["internal_notice"], "（内部资料　　不得外传）")
        self.assertEqual(roles["letter_contact"], "（联系人：张三；电话：12345678）")
        self.assertNotIn("front_matter_order", {finding.code for finding in report.findings})
        self.assertTrue(
            {
                "letter_header_textbox_missing",
                "letter_top_rule_missing",
                "letter_bottom_rule_missing",
                "letter_first_page_number_wrong",
            }.issubset({finding.code for finding in report.findings})
        )

    def test_detects_strict_red_head_and_embedded_regulation_roles(self):
        doc = Document()
        for text in [
            "内部资料不得外传",
            "某某集团有限公司文件",
            "某发〔2026〕54号",
            "关于印发安全生产管理办法的通知",
            "各部门：",
            "现将管理办法印发给你们。",
            "某某集团有限公司",
            "2026年6月22日",
            "KXX/KXXX011-2026",
            "安全生产管理办法",
            "第一章 总则",
            "第一条 为加强管理，制定本办法。",
        ]:
            doc.add_paragraph(text)

        report = OfficialDocumentAuditor().audit_document(doc)
        roles = {element.name: element.text for element in report.elements}

        self.assertEqual(report.profile, "red_head")
        self.assertEqual(roles["internal_notice"], "内部资料不得外传")
        self.assertEqual(roles["regulation_code"], "KXX/KXXX011-2026")
        self.assertEqual(roles["regulation_title"], "安全生产管理办法")
        self.assertIn("regulation_chapter", roles)
        self.assertIn("regulation_article", roles)

    def test_detects_red_head_structure(self):
        doc = Document()
        doc.add_paragraph("某某市人民政府文件")
        doc.add_paragraph("某政发〔2026〕1号")
        doc.add_paragraph("签发人：张三")
        doc.add_paragraph("关于推进办公助手建设的通知")
        doc.add_paragraph("各区人民政府：")
        doc.add_paragraph("现将有关事项通知如下。")
        doc.add_paragraph("某某市人民政府")
        doc.add_paragraph("2026年6月13日")

        report = OfficialDocumentAuditor().audit_document(doc)

        self.assertTrue(report.is_red_head)
        self.assertEqual(report.profile, "red_head")
        self.assertEqual(report.count("error"), 0)
        roles = {element.name: element.text for element in report.elements}
        self.assertEqual(roles["red_head"], "某某市人民政府文件")
        self.assertEqual(roles["document_number"], "某政发〔2026〕1号")
        self.assertEqual(roles["title"], "关于推进办公助手建设的通知")

    def test_missing_title_is_error(self):
        doc = Document()
        doc.add_paragraph("某政发〔2026〕1号")

        report = OfficialDocumentAuditor().audit_document(doc)

        self.assertEqual(report.count("error"), 1)
        self.assertEqual(report.findings[0].code, "missing_title")

    def test_missing_red_head_document_number_warns(self):
        doc = Document()
        doc.add_paragraph("某某市人民政府文件")
        doc.add_paragraph("关于推进办公助手建设的通知")
        doc.add_paragraph("各区人民政府：")
        doc.add_paragraph("正文。")
        doc.add_paragraph("2026年6月13日")

        report = OfficialDocumentAuditor().audit_document(doc)

        self.assertTrue(report.is_red_head)
        self.assertTrue(any(f.code == "missing_document_number" for f in report.findings))

    def test_print_org_date_does_not_override_document_date(self):
        doc = Document()
        doc.add_paragraph("某某市人民政府文件")
        doc.add_paragraph("某政发〔2026〕1号")
        doc.add_paragraph("关于推进办公助手建设的通知")
        doc.add_paragraph("各区人民政府：")
        doc.add_paragraph("正文。")
        doc.add_paragraph("某某市人民政府")
        doc.add_paragraph("2026年6月13日")
        doc.add_paragraph("抄送：市委办公室。")
        doc.add_paragraph("某某市人民政府办公室 2026年6月14日印发")

        report = OfficialDocumentAuditor().audit_document(doc)
        elements = {element.name: element for element in report.elements}

        self.assertEqual(elements["signatory"].block_index, 5)
        self.assertEqual(elements["date"].block_index, 6)
        self.assertEqual(elements["copy_to"].block_index, 7)
        self.assertEqual(elements["print_org_date"].block_index, 8)

    def test_page_layout_finding_is_fixable(self):
        doc = Document()
        doc.add_paragraph("关于推进办公助手建设的通知")
        doc.add_paragraph("正文。")
        config = OfficeToolConfig()
        report = OfficialDocumentAuditor(config).audit_document(doc)

        layout_findings = [f for f in report.findings if f.code.startswith("layout_")]
        self.assertTrue(layout_findings)
        self.assertTrue(all(f.can_fix for f in layout_findings))

    def test_page_layout_and_grid_check_all_sections(self):
        doc = Document()
        doc.add_paragraph("关于测试多节版式的通知")
        second = doc.add_section(WD_SECTION.NEW_PAGE)
        second.top_margin = Cm(1)
        doc.add_paragraph("第二节正文。")

        report = OfficialDocumentAuditor().audit_document(doc)
        codes = {finding.code for finding in report.findings}

        self.assertIn("layout_margin_top_cm_section_2", codes)
        self.assertIn("layout_document_grid_charsPerLine_section_2", codes)

    def test_document_grid_checks_char_space_and_line_pitch(self):
        doc = Document()
        doc.add_paragraph("关于测试文档网格的通知")
        grid = doc.sections[0]._sectPr.find(qn("w:docGrid"))
        if grid is None:
            from docx.oxml import OxmlElement

            grid = OxmlElement("w:docGrid")
            doc.sections[0]._sectPr.append(grid)
        grid.set(qn("w:charsPerLine"), "28")
        grid.set(qn("w:linesPerPage"), "22")
        grid.set(qn("w:charSpace"), "999")
        grid.set(qn("w:linePitch"), "999")

        report = OfficialDocumentAuditor().audit_document(doc)
        codes = {finding.code for finding in report.findings}

        self.assertIn("layout_document_grid_charSpace", codes)
        self.assertIn("layout_document_grid_linePitch", codes)

    def test_unit_rules_flag_date_attachment_and_effective_wording(self):
        doc = Document()
        doc.add_paragraph("关于印发测试办法的通知")
        doc.add_paragraph("各部门：")
        doc.add_paragraph("本办法自发布之日起执行。")
        doc.add_paragraph("附件:测试清单。")
        doc.add_paragraph("办公室")
        doc.add_paragraph("二〇二六年六月十三日")

        config = OfficeToolConfig()
        config.audit.check_title_line_shape = True
        report = OfficialDocumentAuditor(config).audit_document(doc)
        codes = {finding.code for finding in report.findings}

        self.assertIn("date_not_arabic", codes)
        self.assertIn("attachment_colon_not_fullwidth", codes)
        self.assertIn("attachment_name_has_punctuation", codes)
        self.assertIn("bad_effective_date_wording", codes)

    def test_unit_rules_hint_long_title_and_imprint_lines(self):
        doc = Document()
        doc.add_paragraph("关于进一步加强机关内部综合事务协同办理规范化管理和重点工作闭环落实的通知")
        doc.add_paragraph("各部门：")
        doc.add_paragraph("正文。")
        doc.add_paragraph("办公室")
        doc.add_paragraph("2026年6月13日")
        doc.add_paragraph("抄送：市委办公室。")
        doc.add_paragraph("办公室 2026年6月14日印发")

        config = OfficeToolConfig()
        config.audit.check_title_line_shape = True
        report = OfficialDocumentAuditor(config).audit_document(doc)
        codes = {finding.code for finding in report.findings}

        self.assertIn("title_line_shape_hint", codes)
        self.assertIn("imprint_lines_missing", codes)

    def test_front_matter_order_is_checked(self):
        doc = Document()
        doc.add_paragraph("特急")
        doc.add_paragraph("000001")
        doc.add_paragraph("某某市人民政府文件")
        doc.add_paragraph("某政发〔2026〕8号")
        doc.add_paragraph("关于测试版头顺序的通知")
        doc.add_paragraph("2026年6月13日")

        report = OfficialDocumentAuditor().audit_document(doc)
        codes = {finding.code for finding in report.findings}

        self.assertIn("front_matter_order", codes)

    def test_unit_rules_flag_document_number_and_attachment_indent(self):
        doc = Document()
        doc.add_paragraph("某某市人民政府文件")
        doc.add_paragraph("某政发[2026] 8号")
        doc.add_paragraph("关于开展测试工作的通知")
        doc.add_paragraph("各部门：")
        doc.add_paragraph("正文。")
        doc.add_paragraph("附件：测试清单")
        doc.add_paragraph("办公室")
        doc.add_paragraph("2026年6月13日")

        report = OfficialDocumentAuditor().audit_document(doc)
        codes = {finding.code for finding in report.findings}

        self.assertIn("document_number_format_irregular", codes)
        self.assertIn("attachment_note_indent", codes)


if __name__ == "__main__":
    unittest.main()
