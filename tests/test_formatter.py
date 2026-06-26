from pathlib import Path
import re
import sys
import tempfile
import unittest

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, RGBColor

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from office_tool.config import OfficeToolConfig
from office_tool.formatter import OfficialDocumentFormatter
from office_tool.io import load_document


class FormatterTests(unittest.TestCase):
    def test_generation_never_invents_required_business_values(self):
        doc = Document()
        doc.add_paragraph("关于测试工作的通知")
        config = OfficeToolConfig()
        config.audit.profile = "red_head"
        config.generation.add_red_head = True

        with self.assertRaisesRegex(ValueError, "红头名称.*发文字号"):
            OfficialDocumentFormatter(config).format_document(doc)

    def test_export_generates_standard_red_head_and_imprint_through_formatter(self):
        doc = Document()
        doc.add_paragraph("关于推进重点工作的通知")
        doc.add_paragraph("有关单位：")
        doc.add_paragraph("请认真组织实施。")
        config = OfficeToolConfig()
        config.audit.profile = "red_head"
        config.generation.add_red_head = True
        config.generation.add_imprint = True
        config.generation.red_head_title = "某某集团有限公司文件"
        config.generation.document_number = "某发〔2026〕68号"
        config.generation.copy_to = "抄送：有关单位，有关部门"
        config.generation.print_organization = "某某集团办公室"
        config.generation.print_date = "2026年6月23日印发"

        report = OfficialDocumentFormatter(config).format_document(doc)
        texts = [paragraph.text.strip() for paragraph in doc.paragraphs]

        self.assertEqual(texts.count("某某集团有限公司文件"), 1)
        self.assertEqual(texts.count("某发〔2026〕68号"), 1)
        self.assertIn("抄送：有关单位，有关部门。", texts)
        self.assertTrue(any(text.startswith("某某集团办公室") and text.endswith("2026年6月23日印发") for text in texts))
        self.assertEqual(report.stats["generated_content"], ["red_head", "imprint"])
        self.assertIsNotNone(_shape_with_id(doc, "OfficeToolRedSeparator"))
        self.assertIsNotNone(_shape_with_id(doc, "OfficeToolImprintTopLine"))
        imprint_spacers = sum(
            1
            for paragraph in doc.paragraphs
            if paragraph._p.pPr is not None
            and paragraph._p.pPr.pStyle is not None
            and paragraph._p.pPr.pStyle.get(qn("w:val")) == "OfficeToolImprintSpacer"
        )
        self.assertGreater(imprint_spacers, 22)

    def test_export_generates_simple_imprint_when_copy_to_is_blank(self):
        doc = Document()
        doc.add_paragraph("关于推进重点工作的通知")
        doc.add_paragraph("有关单位：")
        doc.add_paragraph("请认真组织实施。")
        config = OfficeToolConfig()
        config.audit.profile = "red_head"
        config.generation.add_imprint = True
        config.generation.print_organization = "某某集团有限公司办公室"
        config.generation.print_date = "2026年6月23日"

        report = OfficialDocumentFormatter(config).format_document(doc)
        texts = [paragraph.text.strip() for paragraph in doc.paragraphs]

        self.assertEqual(report.stats["generated_content"], ["simple_imprint"])
        self.assertFalse(any(text.startswith("抄送：") for text in texts))
        self.assertTrue(any(text.startswith("某某集团有限公司办公室") and text.endswith("2026年6月23日") for text in texts))
        self.assertIsNotNone(_shape_with_id(doc, "OfficeToolSimpleImprintTopLine"))
        self.assertIsNotNone(_shape_with_id(doc, "OfficeToolSimpleImprintBottomLine"))

    def test_generation_does_not_duplicate_partial_existing_red_head(self):
        doc = Document()
        doc.add_paragraph("某某集团有限公司文件")
        doc.add_paragraph("关于现有红头的通知")
        config = OfficeToolConfig()
        config.audit.profile = "red_head"
        config.generation.add_red_head = True
        config.generation.red_head_title = "另一单位文件"
        config.generation.document_number = "另发〔2026〕1号"

        report = OfficialDocumentFormatter(config).format_document(doc)
        texts = [paragraph.text.strip() for paragraph in doc.paragraphs]

        self.assertNotIn("另一单位文件", texts)
        self.assertIn("另发〔2026〕1号", texts)
        self.assertEqual(report.stats["generated_content"], ["red_head"])

    def test_export_generates_meeting_header_and_distribution(self):
        doc = Document()
        doc.add_paragraph("关于研究重点工作的会议纪要")
        doc.add_paragraph("会议研究了有关事项。")
        config = OfficeToolConfig()
        config.audit.profile = "meeting_minutes"
        config.generation.add_red_head = True
        config.generation.add_imprint = True
        config.generation.meeting_number = "12"
        config.generation.meeting_organization = "某某集团办公室"
        config.generation.meeting_date = "2026年6月23日"
        config.generation.distribution = "分送：有关单位，有关部门"

        report = OfficialDocumentFormatter(config).format_document(doc)
        texts = [paragraph.text.strip() for paragraph in doc.paragraphs]

        self.assertIn("会 议 纪 要", texts)
        self.assertIn("（ 12 ）", texts)
        self.assertTrue(any(text.startswith("某某集团办公室") and text.endswith("2026年6月23日") for text in texts))
        self.assertIn("分送：有关单位，有关部门。", texts)
        self.assertEqual(report.stats["generated_content"], ["red_head", "distribution"])

    def test_export_generates_letter_header_without_imprint(self):
        doc = Document()
        doc.add_paragraph("关于商请协助办理事项的函")
        doc.add_paragraph("有关单位：")
        doc.add_paragraph("请协助办理。")
        config = OfficeToolConfig()
        config.audit.profile = "letter_head"
        config.generation.add_red_head = True
        config.generation.add_imprint = True
        config.generation.red_head_title = "某某集团有限公司"
        config.generation.document_number = "某部门函〔2026〕42号"

        report = OfficialDocumentFormatter(config).format_document(doc)
        texts = [paragraph.text.strip() for paragraph in doc.paragraphs]

        text_box = next(
            (shape for shape in doc._element.body.iter("{urn:schemas-microsoft-com:vml}shape") if shape.get("id") == "OfficeToolLetterHeadTextBox"),
            None,
        )
        self.assertIsNotNone(text_box)
        self.assertEqual("".join(node.text or "" for node in text_box.iter(qn("w:t"))), "某某集团有限公司")
        self.assertIn("某部门函〔2026〕42号", texts)
        self.assertIn("（内部资料　　不得外传）", texts)
        self.assertEqual(report.stats["generated_content"], ["red_head"])
        self.assertFalse(any(text.startswith("抄送：") for text in texts))

    def test_export_completes_letter_header_when_document_number_exists(self):
        doc = Document()
        doc.add_paragraph("某部门函〔2026〕42号")
        doc.add_paragraph("关于商请协助办理事项的函")
        doc.add_paragraph("有关单位：")
        doc.add_paragraph("请协助办理。")
        config = OfficeToolConfig()
        config.audit.profile = "letter_head"
        config.generation.add_red_head = True
        config.generation.red_head_title = "某某集团有限公司"

        report = OfficialDocumentFormatter(config).format_document(doc)
        texts = [paragraph.text.strip() for paragraph in doc.paragraphs]

        text_box = next(
            (shape for shape in doc._element.body.iter("{urn:schemas-microsoft-com:vml}shape") if shape.get("id") == "OfficeToolLetterHeadTextBox"),
            None,
        )
        self.assertIsNotNone(text_box)
        self.assertEqual(texts.count("某部门函〔2026〕42号"), 1)
        self.assertEqual("".join(node.text or "" for node in text_box.iter(qn("w:t"))), "某某集团有限公司")
        self.assertEqual(report.stats["generated_content"], ["red_head"])
        self.assertFalse(any(finding.code.startswith("letter_") for finding in report.findings))

    def test_distribution_moves_to_even_page_for_standard_document(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_distribution_") as tmp:
            source = Path(tmp) / "distribution.docx"
            output = Path(tmp) / "distribution_out.docx"
            doc = Document()
            for text in [
                "关于测试分送版记的通知",
                "正文内容。",
                "某某单位办公室",
                "2026年6月22日",
                "分送：有关部门，有关单位",
            ]:
                doc.add_paragraph(text)
            doc.save(source)

            report = OfficialDocumentFormatter().format_path(source, output)
            formatted = Document(output)

            self.assertEqual(report.profile, "standard")
            self.assertNotIn("distribution_layout_missing", {finding.code for finding in report.findings})
            self.assertIsNotNone(_shape_with_id(formatted, "OfficeToolDistributionTopLine"))
            self.assertIsNotNone(_shape_with_id(formatted, "OfficeToolDistributionBottomLine"))
            self.assertFalse(any(br.get(qn("w:type")) == "page" for br in formatted._element.body.iter(qn("w:br"))))

    def test_distribution_spacing_accounts_for_tables_and_page_break_before(self):
        def formatted_spacer_count(path: Path, with_table: bool, with_page_break: bool) -> int:
            doc = Document()
            doc.add_paragraph("关于测试分送版记的通知")
            doc.add_paragraph("正文内容。")
            if with_table:
                table = doc.add_table(rows=5, cols=2)
                for row_index, row in enumerate(table.rows, start=1):
                    row.cells[0].text = f"事项{row_index}"
                    row.cells[1].text = "需要处理的内容"
            anchor = doc.add_paragraph("分页后的正文。")
            anchor.paragraph_format.page_break_before = with_page_break
            doc.add_paragraph("分送：有关部门，有关单位")
            OfficialDocumentFormatter().format_document(doc)
            doc.save(path)
            return sum(
                1
                for paragraph in doc.paragraphs
                if paragraph._p.pPr is not None
                and paragraph._p.pPr.pStyle is not None
                and paragraph._p.pPr.pStyle.get(qn("w:val")) == "OfficeToolDistributionSpacer"
            )

        with tempfile.TemporaryDirectory(prefix="office_tool_distribution_blocks_") as tmp:
            root = Path(tmp)
            plain = formatted_spacer_count(root / "plain.docx", False, False)
            with_table = formatted_spacer_count(root / "table.docx", True, False)
            with_page_break = formatted_spacer_count(root / "break.docx", False, True)

        self.assertLess(with_table, plain)
        self.assertLess(with_page_break, plain)

    def test_format_meeting_minutes_and_even_page_distribution(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_meeting_minutes_") as tmp:
            source = Path(tmp) / "meeting.docx"
            output = Path(tmp) / "meeting_out.docx"
            doc = Document()
            for text in [
                "内部资料不得外传",
                "会议纪要",
                "（12）",
                "某某单位办公室    2026年6月22日",
                "关于研究重点工作的会议纪要",
                "会议研究了有关事项。",
                "出席：张三、李四、王五、赵六、钱七、孙八",
                "分送：有关部门，有关单位",
            ]:
                doc.add_paragraph(text)
            doc.save(source)

            report = OfficialDocumentFormatter().format_path(source, output)
            formatted = Document(output)
            internal = _paragraph_with_text(formatted, "内部资料\n不得外传")
            red_head = _paragraph_with_text(formatted, "会 议 纪 要")
            number = _paragraph_with_text(formatted, "（ 12 ）")
            issue = _paragraph_with_prefix(formatted, "某某单位办公室")
            attendees = _paragraph_with_prefix(formatted, "出席：")
            distribution = _paragraph_with_prefix(formatted, "分送：")
            vml_ns = "urn:schemas-microsoft-com:vml"
            red_line = _shape_with_id(formatted, "OfficeToolRedSeparator")
            top_line = _shape_with_id(formatted, "OfficeToolDistributionTopLine")
            bottom_line = _shape_with_id(formatted, "OfficeToolDistributionBottomLine")

            self.assertEqual(report.profile, "meeting_minutes")
            self.assertFalse(any(f.code.startswith("meeting_") or f.code.startswith("distribution_") for f in report.findings))
            self.assertEqual(_blank_count_between(formatted, internal, red_head), 1)
            blank = formatted.paragraphs[_paragraph_index(formatted, internal) + 1]
            self.assertEqual(blank.paragraph_format.line_spacing.pt, 38)
            self.assertEqual(_east_asia_font(red_head.runs[0]), "华文中宋")
            self.assertEqual(red_head.runs[0].font.size.pt, 49)
            self.assertEqual(red_head.paragraph_format.line_spacing.pt, 48)
            self.assertEqual(number.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(issue.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(_indent_chars(issue, "leftChars"), 100)
            self.assertEqual(_indent_chars(issue, "rightChars"), 100)
            self.assertLessEqual(max((len(part) for part in re.findall(r" +", issue.text)), default=0), 24)
            self.assertIsNotNone(red_line)
            self.assertEqual(red_line.get("strokeweight"), "1.4pt")
            self.assertEqual(attendees.runs[0].text, "出席：")
            self.assertEqual(_east_asia_font(attendees.runs[0]), "黑体")
            self.assertEqual(_east_asia_font(attendees.runs[1]), "仿宋_GB2312")
            self.assertEqual(_indent_chars(attendees, "leftChars"), 200)
            self.assertEqual(_indent_chars(attendees, "hangingChars"), 310)
            self.assertEqual(distribution.text, "分送：有关部门，有关单位。")
            self.assertEqual(top_line.get("strokeweight"), "1pt")
            self.assertEqual(bottom_line.get("strokeweight"), "1pt")
            self.assertEqual({top_line.get("to"), bottom_line.get("to")}, {"442.20pt,0"})
            spacer_count = sum(
                1
                for paragraph in formatted.paragraphs
                if paragraph._p.pPr is not None
                and paragraph._p.pPr.pStyle is not None
                and paragraph._p.pPr.pStyle.get(qn("w:val")) == "OfficeToolDistributionSpacer"
            )
            self.assertGreater(spacer_count, 0)
            self.assertFalse(any(br.get(qn("w:type")) == "page" for br in formatted._element.body.iter(qn("w:br"))))

    def test_format_letter_head_builds_first_page_furniture(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_letter_head_") as tmp:
            source = Path(tmp) / "letter.docx"
            output = Path(tmp) / "letter_out.docx"
            doc = Document()
            for text in [
                "某某集团有限公司",
                "某部门函〔2026〕42号",
                "关于商请协助办理有关事项的函",
                "有关单位：",
                "请协助办理有关事项。",
                "某某集团有限公司",
                "2026年6月22日",
                "（内部资料　　不得外传）",
                "（联系人：张三；电话：12345678）",
            ]:
                doc.add_paragraph(text)
            doc.save(source)

            report = OfficialDocumentFormatter().format_path(source, output)
            formatted = Document(output)
            section = formatted.sections[0]
            document_number = _paragraph_with_text(formatted, "某部门函〔2026〕42号")
            title = _paragraph_with_text(formatted, "关于商请协助办理有关事项的函")
            internal_notice = _paragraph_with_text(formatted, "（内部资料　　不得外传）")
            contact = _paragraph_with_text(formatted, "（联系人：张三；电话：12345678）")
            vml_ns = "urn:schemas-microsoft-com:vml"
            body_shapes = list(formatted._element.body.iter(f"{{{vml_ns}}}shape"))
            text_box = next(shape for shape in body_shapes if shape.get("id") == "OfficeToolLetterHeadTextBox")
            text_box_text = "".join(node.text or "" for node in text_box.iter(qn("w:t")))
            top_line = _shape_with_id(formatted, "OfficeToolLetterRedTop")
            bottom_line = _shape_with_id(formatted, "OfficeToolLetterRedBottom")

            self.assertEqual(report.profile, "letter_head")
            self.assertFalse(any(finding.code.startswith("letter_") for finding in report.findings))
            self.assertTrue(section.different_first_page_header_footer)
            self.assertAlmostEqual(section.bottom_margin.cm, 2.5, places=2)
            self.assertEqual(text_box_text, "某某集团有限公司")
            self.assertEqual(top_line.get("strokeweight"), "4pt")
            self.assertEqual(top_line.find(f"{{{vml_ns}}}stroke").get("linestyle"), "thickThin")
            self.assertEqual(bottom_line.get("strokeweight"), "4pt")
            self.assertEqual(bottom_line.find(f"{{{vml_ns}}}stroke").get("linestyle"), "thinThick")
            self.assertNotIn("PAGE", "".join(section.first_page_footer._element.itertext()))
            self.assertIn("PAGE", "".join(section.footer._element.itertext()))
            self.assertEqual(document_number.alignment, WD_ALIGN_PARAGRAPH.RIGHT)
            self.assertEqual(_blank_count_after(formatted, document_number), 2)
            self.assertEqual(title.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(title.runs[0].font.size.pt, 22)
            self.assertEqual(_east_asia_font(internal_notice.runs[0]), "黑体")
            self.assertEqual(internal_notice.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
            self.assertEqual(_east_asia_font(contact.runs[0]), "仿宋_GB2312")
            self.assertEqual(contact.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

    def test_format_red_head_keeps_title_separate(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_test_") as tmp:
            source = Path(tmp) / "red.docx"
            output = Path(tmp) / "red_out.docx"
            doc = Document()
            doc.add_paragraph("某某市人民政府文件")
            doc.add_paragraph("某政发〔2026〕1号")
            doc.add_paragraph("关于推进办公助手建设的通知")
            doc.add_paragraph("各区人民政府：")
            doc.add_paragraph("一、总体要求")
            doc.add_paragraph("正文内容。")
            doc.add_paragraph("某某市人民政府")
            doc.add_paragraph("2026年6月13日")
            doc.save(source)

            report = OfficialDocumentFormatter().format_path(source, output)
            formatted = Document(output)
            red_head = _paragraph_with_text(formatted, "某某市人民政府文件")
            title = _paragraph_with_text(formatted, "关于推进办公助手建设的通知")

            self.assertTrue(report.is_red_head)
            self.assertEqual(red_head.runs[0].font.color.rgb, RGBColor(0xFF, 0, 0))
            self.assertEqual(_east_asia_font(red_head.runs[0]), "华文中宋")
            self.assertEqual(red_head.runs[0].font.size.pt, 42)
            self.assertEqual(title.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(title.runs[0].font.size.pt, 22)
            self.assertAlmostEqual(formatted.sections[0].top_margin.cm, 3.7, places=2)
            self.assertAlmostEqual(formatted.sections[0].right_margin.cm, 2.6, places=2)

    def test_format_report_reflects_final_page_setup(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_final_report_") as tmp:
            source = Path(tmp) / "sample.docx"
            output = Path(tmp) / "sample_out.docx"
            doc = Document()
            doc.sections[0].top_margin = Cm(1)
            doc.add_paragraph("关于测试的通知")
            doc.add_paragraph("各部门：")
            doc.add_paragraph("正文。")
            doc.add_section(WD_SECTION.NEW_PAGE)
            doc.add_paragraph("第二节正文。")
            doc.add_paragraph("办公室")
            doc.add_paragraph("2026年6月13日")
            doc.save(source)

            report = OfficialDocumentFormatter().format_path(source, output)

            self.assertFalse(any(f.code.startswith("layout_") for f in report.findings))

    def test_format_writes_document_grid_and_odd_even_page_numbers(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_grid_") as tmp:
            source = Path(tmp) / "sample.docx"
            output = Path(tmp) / "sample_out.docx"
            doc = Document()
            doc.add_paragraph("关于测试的通知")
            doc.add_paragraph("各部门：")
            doc.add_paragraph("正文。")
            doc.add_section(WD_SECTION.NEW_PAGE)
            doc.add_paragraph("第二节正文。")
            doc.add_paragraph("办公室")
            doc.add_paragraph("2026年6月13日")
            doc.save(source)

            OfficialDocumentFormatter().format_path(source, output)
            formatted = Document(output)
            self.assertAlmostEqual(formatted.sections[0].page_width.cm, 21.0, places=2)
            self.assertAlmostEqual(formatted.sections[0].page_height.cm, 29.7, places=2)
            self.assertAlmostEqual(formatted.sections[1].page_width.cm, 21.0, places=2)
            self.assertAlmostEqual(formatted.sections[1].page_height.cm, 29.7, places=2)
            grid = formatted.sections[0]._sectPr.find(qn("w:docGrid"))
            self.assertIsNotNone(grid)
            self.assertEqual(grid.get(qn("w:type")), "lines")
            self.assertEqual(grid.get(qn("w:charsPerLine")), "28")
            self.assertEqual(grid.get(qn("w:linesPerPage")), "22")
            self.assertEqual(grid.get(qn("w:charSpace")), "-834")
            self.assertEqual(grid.get(qn("w:linePitch")), "579")
            second_grid = formatted.sections[1]._sectPr.find(qn("w:docGrid"))
            self.assertIsNotNone(second_grid)
            self.assertEqual(second_grid.get(qn("w:type")), "lines")
            self.assertEqual(second_grid.get(qn("w:charsPerLine")), "28")
            self.assertEqual(second_grid.get(qn("w:linesPerPage")), "22")
            self.assertEqual(second_grid.get(qn("w:charSpace")), "-834")
            self.assertEqual(second_grid.get(qn("w:linePitch")), "579")
            self.assertGreaterEqual(
                formatted.sections[0].page_width.twips
                - formatted.sections[0].left_margin.twips
                - formatted.sections[0].right_margin.twips,
                round(28 * 15.8 * 20),
            )
            self.assertAlmostEqual(formatted.sections[0].left_margin.cm, 2.8, places=2)
            self.assertAlmostEqual(formatted.sections[0].right_margin.cm, 2.6, places=2)
            self.assertAlmostEqual(formatted.sections[1].left_margin.cm, 2.8, places=2)
            self.assertAlmostEqual(formatted.sections[1].right_margin.cm, 2.6, places=2)
            self.assertEqual(formatted.sections[0].footer.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.RIGHT)
            self.assertEqual(formatted.sections[0].even_page_footer.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertEqual(_indent_chars(formatted.sections[0].footer.paragraphs[0], "rightChars"), 100)
            self.assertEqual(_indent_chars(formatted.sections[0].even_page_footer.paragraphs[0], "leftChars"), 100)
            self.assertIn("— ", formatted.sections[0].footer.paragraphs[0].text)
            self.assertIn(" —", formatted.sections[0].footer.paragraphs[0].text)

    def test_format_fixes_attachment_indent_and_keeps_plain_article_as_body(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_article_") as tmp:
            source = Path(tmp) / "regulation.docx"
            output = Path(tmp) / "regulation_out.docx"
            doc = Document()
            doc.add_paragraph("安全生产管理办法")
            doc.add_paragraph("第一条 为加强管理，制定本办法。")
            doc.add_paragraph("附件：职责清单")
            doc.add_paragraph("办公室")
            doc.add_paragraph("2026年6月13日")
            doc.save(source)

            OfficialDocumentFormatter().format_path(source, output)
            formatted = Document(output)

            attachment = _paragraph_with_text(formatted, "附件：职责清单")
            article = _paragraph_with_prefix(formatted, "第一条")
            self.assertEqual(_indent_chars(attachment, "leftChars"), 200)
            self.assertEqual(_indent_chars(attachment, "hangingChars"), 450)
            self.assertEqual(attachment.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
            self.assertEqual(_east_asia_font(article.runs[0]), "仿宋_GB2312")
            self.assertEqual(_indent_chars(article, "firstLineChars"), 200)

    def test_format_normalizes_punctuation_spaces_and_heading_periods(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_normalize_") as tmp:
            source = Path(tmp) / "normalize.docx"
            output = Path(tmp) / "normalize_out.docx"
            doc = Document()
            doc.add_paragraph("关于 测试 的通知")
            doc.add_paragraph("各科室,车间:")
            doc.add_paragraph("（一）基本 情况")
            doc.add_paragraph("1. 具体 措施")
            doc.add_paragraph("正文 内容:请 按要求 办理!")
            doc.save(source)

            OfficialDocumentFormatter().format_path(source, output)
            formatted = Document(output)

            self.assertEqual(formatted.paragraphs[0].text, "关于测试的通知")
            self.assertEqual(formatted.paragraphs[1].text, "")
            self.assertEqual(formatted.paragraphs[2].text, "各科室，车间：")
            self.assertEqual(formatted.paragraphs[3].text, "（一）基本情况。")
            self.assertEqual(formatted.paragraphs[4].text, "1．具体措施。")
            self.assertEqual(formatted.paragraphs[5].text, "正文内容：请按要求办理！")

    def test_format_body_uses_single_spacing_and_signatory_right_aligns(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_spacing_signatory_") as tmp:
            source = Path(tmp) / "spacing.docx"
            output = Path(tmp) / "spacing_out.docx"
            doc = Document()
            doc.add_paragraph("关于测试的通知")
            doc.add_paragraph("")
            doc.add_paragraph("")
            doc.add_paragraph("")
            doc.add_paragraph("各单位：")
            doc.add_paragraph("1. 正文内容需要单倍行距。")
            doc.add_paragraph("")
            doc.add_paragraph("办公室")
            doc.add_paragraph("2026年6月16日")
            doc.save(source)

            OfficialDocumentFormatter().format_path(source, output)
            formatted = Document(output)
            texts = [paragraph.text.rstrip() for paragraph in formatted.paragraphs]
            body_spacing = formatted.paragraphs[3]._p.pPr.spacing

            self.assertEqual(body_spacing.get(qn("w:line")), "240")
            self.assertEqual(body_spacing.get(qn("w:lineRule")), "auto")
            self.assertEqual(texts[:8], ["关于测试的通知", "", "各单位：", "1．正文内容需要单倍行距。", "", "", "办公室", "2026年6月16日"])
            self.assertEqual(formatted.paragraphs[6].alignment, WD_ALIGN_PARAGRAPH.RIGHT)
            self.assertEqual(formatted.paragraphs[7].alignment, WD_ALIGN_PARAGRAPH.RIGHT)
            title_spacing = formatted.paragraphs[0]._p.pPr.spacing
            signatory_spacing = formatted.paragraphs[6]._p.pPr.spacing
            self.assertEqual(title_spacing.get(qn("w:after")), "0")
            self.assertEqual(signatory_spacing.get(qn("w:before")), "0")
            self.assertEqual(_indent_chars(formatted.paragraphs[6], "rightChars"), None)
            self.assertEqual(len(formatted.paragraphs[6].text) - len(formatted.paragraphs[6].text.rstrip()), 22)
            spacer_spacing = formatted.paragraphs[6].runs[-1]._element.rPr.find(qn("w:spacing"))
            self.assertIsNotNone(spacer_spacing)
            self.assertEqual(_indent_chars(formatted.paragraphs[7], "rightChars"), 400)

    def test_format_centers_long_signatory_against_date(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_long_signatory_") as tmp:
            source = Path(tmp) / "long_signatory.docx"
            output = Path(tmp) / "long_signatory_out.docx"
            doc = Document()
            doc.add_paragraph("关于测试的通知")
            doc.add_paragraph("各单位：")
            doc.add_paragraph("正文内容。")
            doc.add_paragraph("某某单位办公室")
            doc.add_paragraph("2026年6月19日")
            doc.save(source)

            OfficialDocumentFormatter().format_path(source, output)
            formatted = Document(output)
            texts = [paragraph.text.rstrip() for paragraph in formatted.paragraphs]
            signatory_index = texts.index("某某单位办公室")
            date_index = texts.index("2026年6月19日")

            self.assertEqual(_indent_chars(formatted.paragraphs[signatory_index], "rightChars"), None)
            self.assertEqual(
                len(formatted.paragraphs[signatory_index].text)
                - len(formatted.paragraphs[signatory_index].text.rstrip()),
                14,
            )
            self.assertEqual(_indent_chars(formatted.paragraphs[date_index], "rightChars"), 400)

    def test_format_strict_red_head_header_and_floating_separator(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_strict_red_head_") as tmp:
            source = Path(tmp) / "strict_red_head.docx"
            output = Path(tmp) / "strict_red_head_out.docx"
            doc = Document()
            doc.add_paragraph("内部资料不得外传")
            doc.add_paragraph("某某集团有限公司委员会文件")
            doc.add_paragraph("某政发〔2026〕54号")
            doc.add_paragraph("关于开展专项检查的通知")
            doc.add_paragraph("各单位：")
            doc.add_paragraph("请认真组织实施。")
            doc.add_paragraph("某某市人民政府办公室")
            doc.add_paragraph("2026年6月22日")
            doc.save(source)

            report = OfficialDocumentFormatter().format_path(source, output)
            formatted = Document(output)
            internal = _paragraph_with_text(formatted, "内部资料\n不得外传")
            red_head = _paragraph_with_text(formatted, "某某集团有限公司委员会文件")
            number = _paragraph_with_text(formatted, "某政发〔2026〕54号")
            title = _paragraph_with_text(formatted, "关于开展专项检查的通知")
            line = next(p for p in formatted.paragraphs if _floating_line(p) is not None)
            line_element = _floating_line(line)

            self.assertEqual(report.profile, "red_head")
            self.assertEqual(internal.alignment, WD_ALIGN_PARAGRAPH.RIGHT)
            self.assertEqual(_east_asia_font(internal.runs[0]), "黑体")
            self.assertEqual(internal.runs[0].font.size.pt, 14)
            self.assertEqual(_blank_count_between(formatted, internal, red_head), 1)
            self.assertEqual(_blank_count_between(formatted, red_head, number), 2)
            self.assertEqual(red_head.paragraph_format.line_spacing.pt, 48)
            self.assertLess(red_head.runs[0].font.size.pt, 42)
            self.assertGreaterEqual(red_head.runs[0].font.size.pt, 22)
            self.assertGreaterEqual(_run_spacing(red_head.runs[0]), -20)
            self.assertEqual(number.paragraph_format.line_spacing.pt, 30)
            self.assertEqual(_east_asia_font(number.runs[0]), "仿宋_GB2312")
            self.assertEqual(line_element.get("strokeweight"), "1.4pt")
            self.assertEqual(line_element.get("strokecolor"), "#FF0000")
            self.assertIn("442.20pt", line_element.get("to"))
            self.assertEqual(_blank_count_between(formatted, line, title), 2)

    def test_format_red_head_with_regulation_and_bottom_imprint(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_red_regulation_") as tmp:
            source = Path(tmp) / "red_regulation.docx"
            output = Path(tmp) / "red_regulation_out.docx"
            doc = Document()
            for text in [
                "内部资料不得外传",
                "某某集团有限公司文件",
                "某发〔2026〕54号",
                "关于印发安全生产管理办法的通知",
                "各部门：",
                "现将管理办法印发给你们，请认真执行。",
                "某某集团有限公司",
                "2026年6月22日",
                "KXX/KXXX011-2026",
                "安全生产管理办法",
                "第一章 总则",
                "第一条 为加强安全生产管理，制定本办法。",
                "第二条 本办法适用于所属各单位。",
                "第二章 管理职责",
                "第三条 各部门应当落实管理职责。",
                "抄送:上级单位、所属各单位",
                "某某集团有限公司办公室 2026年6月22日印发",
            ]:
                doc.add_paragraph(text)
            doc.save(source)

            report = OfficialDocumentFormatter().format_path(source, output)
            formatted = Document(output)
            code = _paragraph_with_text(formatted, "KXX/KXXX011-2026")
            regulation_title = _paragraph_with_text(formatted, "安全生产管理办法")
            chapter_one = _paragraph_with_text(formatted, "第一章　总则")
            chapter_two = _paragraph_with_text(formatted, "第二章　管理职责")
            article = _paragraph_with_prefix(formatted, "第一条　")
            copy_to = _paragraph_with_prefix(formatted, "抄送：")
            print_line = _paragraph_with_prefix(formatted, "某某集团有限公司办公室")

            self.assertEqual(report.profile, "red_head")
            self.assertNotIn("imprint_lines_missing", {finding.code for finding in report.findings})
            self.assertIsNone(code.paragraph_format.page_break_before)
            previous = code._p.getprevious()
            self.assertTrue(any(br.get(qn("w:type")) == "page" for br in previous.iter(qn("w:br"))))
            self.assertEqual(_east_asia_font(code.runs[0]), "Times New Roman")
            self.assertEqual(code.runs[0].font.size.pt, 16)
            self.assertEqual(code.paragraph_format.line_spacing.pt, 30)
            self.assertEqual(_blank_count_between(formatted, code, regulation_title), 1)
            self.assertEqual(_east_asia_font(regulation_title.runs[0]), "华文中宋")
            self.assertEqual(_blank_count_between(formatted, regulation_title, chapter_one), 1)
            self.assertEqual(_east_asia_font(chapter_one.runs[0]), "黑体")
            self.assertEqual(chapter_one.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(_blank_count_before(formatted, chapter_two), 1)
            self.assertEqual(_blank_count_after(formatted, chapter_two), 1)
            self.assertEqual(article.runs[0].text, "第一条　")
            self.assertEqual(_east_asia_font(article.runs[0]), "黑体")
            self.assertEqual(_east_asia_font(article.runs[1]), "仿宋_GB2312")
            self.assertEqual(copy_to.text, "抄送：上级单位，所属各单位。")
            self.assertEqual(_indent_chars(copy_to, "leftChars"), 100)
            self.assertEqual(_indent_chars(copy_to, "rightChars"), 100)
            self.assertGreaterEqual(_internal_spaces(print_line.text), 4)
            self.assertTrue(print_line.text.endswith("2026年6月22日印发"))
            top_line = _shape_with_id(formatted, "OfficeToolImprintTopLine")
            middle_line = _shape_with_id(formatted, "OfficeToolImprintMiddleLine")
            bottom_line = _shape_with_id(formatted, "OfficeToolImprintBottomLine")
            self.assertEqual(top_line.get("strokeweight"), "1pt")
            self.assertEqual(middle_line.get("strokeweight"), "0.6pt")
            self.assertEqual(bottom_line.get("strokeweight"), "1pt")
            self.assertEqual({top_line.get("to"), middle_line.get("to"), bottom_line.get("to")}, {"442.20pt,0"})
            line_paragraphs = [_paragraph_containing_shape(formatted, line) for line in (top_line, middle_line, bottom_line)]
            self.assertEqual(len({id(paragraph._p) for paragraph in line_paragraphs}), 3)
            imprint_paragraphs = [line_paragraphs[0], copy_to, line_paragraphs[1], print_line, line_paragraphs[2]]
            self.assertEqual([_paragraph_index(formatted, p) for p in imprint_paragraphs], sorted(_paragraph_index(formatted, p) for p in imprint_paragraphs))
            self.assertTrue(all(_frame_value(paragraph, "y") is None for paragraph in imprint_paragraphs))
            self.assertTrue(
                all(
                    paragraph._p.pPr.find(qn(tag)) is None
                    for paragraph in imprint_paragraphs
                    for tag in ("w:keepNext", "w:keepLines", "w:pageBreakBefore")
                )
            )

            for paragraph in formatted.paragraphs:
                spacing = paragraph._p.pPr.find(qn("w:spacing"))
                self.assertIn(spacing.get(qn("w:lineRule")), {"auto", "exact"})
                if spacing.get(qn("w:lineRule")) == "auto":
                    self.assertEqual(spacing.get(qn("w:line")), "240")

    def test_format_keeps_title_date_with_title_block(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_title_date_") as tmp:
            source = Path(tmp) / "title_date.docx"
            output = Path(tmp) / "title_date_out.docx"
            doc = Document()
            doc.add_paragraph("这是工作总结标题")
            doc.add_paragraph("（2026年1月1日-2026年6月30日）")
            doc.add_paragraph("各单位：")
            doc.add_paragraph("正文内容。")
            doc.save(source)

            OfficialDocumentFormatter().format_path(source, output)
            formatted = Document(output)

            self.assertEqual(formatted.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(formatted.paragraphs[1].alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(_east_asia_font(formatted.paragraphs[0].runs[0]), "华文中宋")
            self.assertEqual(_east_asia_font(formatted.paragraphs[1].runs[0]), "仿宋_GB2312")
            self.assertEqual(_indent_chars(formatted.paragraphs[1], "firstLineChars"), None)
            self.assertEqual(formatted.paragraphs[2].text, "")
            self.assertEqual(formatted.paragraphs[3].text, "各单位：")

    def test_format_single_date_below_title_uses_centered_body_font(self):
        doc = Document()
        doc.add_paragraph("2026年上半年工作总结")
        doc.add_paragraph("2026年6月23日")
        doc.add_paragraph("各单位：")
        doc.add_paragraph("正文内容。")

        OfficialDocumentFormatter().format_document(doc)
        date_line = _paragraph_with_text(doc, "2026年6月23日")

        self.assertEqual(date_line.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(_east_asia_font(date_line.runs[0]), "Times New Roman")
        self.assertEqual(_east_asia_font(date_line.runs[1]), "仿宋_GB2312")
        self.assertEqual(date_line.runs[0].font.size.pt, 16)
        self.assertEqual(_indent_chars(date_line, "firstLineChars"), None)

    def test_body_date_inside_sentence_is_not_formatted_as_document_date(self):
        doc = Document()
        doc.add_paragraph("关于推进重点工作的通知")
        paragraph = doc.add_paragraph("各单位要在2026年5月5日前完成资料报送。")

        report = OfficialDocumentFormatter().format_document(doc)

        self.assertNotIn("date", {element.name for element in report.elements})
        self.assertNotEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.RIGHT)
        self.assertEqual(_indent_chars(paragraph, "rightChars"), None)

    def test_h3_heading_digits_are_split_to_times_new_roman(self):
        doc = Document()
        doc.add_paragraph("关于推进重点工作的通知")
        paragraph = doc.add_paragraph("1.开展专项检查。")

        OfficialDocumentFormatter().format_document(doc)

        digit_run = next(run for run in paragraph.runs if run.text == "1")
        self.assertEqual(_east_asia_font(digit_run), "Times New Roman")
        self.assertEqual(_ascii_font(digit_run), "Times New Roman")

    def test_format_handles_split_title_lines(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_split_title_") as tmp:
            source = Path(tmp) / "split_title.docx"
            output = Path(tmp) / "split_title_out.docx"
            doc = Document()
            doc.add_paragraph("某某单位")
            doc.add_paragraph("关于规范公文处理有关事项")
            doc.add_paragraph("的通知")
            doc.add_paragraph("各单位：")
            doc.add_paragraph("正文内容。")
            doc.save(source)

            OfficialDocumentFormatter().format_path(source, output)
            formatted = Document(output)

            for index in (0, 1, 2):
                self.assertEqual(formatted.paragraphs[index].alignment, WD_ALIGN_PARAGRAPH.CENTER)
                self.assertEqual(_east_asia_font(formatted.paragraphs[index].runs[0]), "华文中宋")
            self.assertEqual(formatted.paragraphs[3].text, "")
            self.assertEqual(formatted.paragraphs[4].text, "各单位：")

    def test_format_h2_inline_body_only_styles_heading_sentence(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_h2_inline_") as tmp:
            source = Path(tmp) / "h2_inline.docx"
            output = Path(tmp) / "h2_inline_out.docx"
            doc = Document()
            doc.add_paragraph("关于测试的通知")
            doc.add_paragraph("各单位：")
            doc.add_paragraph("（一）这是标题。这是正文内容。")
            doc.save(source)

            OfficialDocumentFormatter().format_path(source, output)
            formatted = Document(output)
            paragraph = next(p for p in formatted.paragraphs if p.text.startswith("（一）"))

            self.assertEqual(paragraph.runs[0].text, "（一）这是标题。")
            self.assertEqual(paragraph.runs[1].text, "这是正文内容。")
            self.assertEqual(_east_asia_font(paragraph.runs[0]), "楷体_GB2312")
            self.assertEqual(_east_asia_font(paragraph.runs[1]), "仿宋_GB2312")

    def test_format_attachment_items_use_fullwidth_dot_and_no_period(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_attachment_items_") as tmp:
            source = Path(tmp) / "attachments.docx"
            output = Path(tmp) / "attachments_out.docx"
            doc = Document()
            doc.add_paragraph("关于测试的通知")
            doc.add_paragraph("各单位：")
            doc.add_paragraph("正文内容。")
            doc.add_paragraph("正文结束段落。")
            doc.add_paragraph("附件：1.关于公文质量检查项目以及整改工作责任分工安排的详细说明材料。")
            doc.add_paragraph("2．关于常见格式问题识别方法和具体纠正措施的完整参考清单。")
            doc.add_paragraph("3．关于年度重点任务推进情况以及阶段目标完成情况的统计汇总表。")
            doc.add_paragraph("4．关于各部门协同办理事项责任人员和时间节点安排的工作台账。")
            doc.add_paragraph("5．关于现场检查发现问题整改闭环及复核确认情况的报告材料。")
            doc.save(source)

            OfficialDocumentFormatter().format_path(source, output)
            formatted = Document(output)
            texts = [p.text for p in formatted.paragraphs]

            first = next(p for p in formatted.paragraphs if p.text.startswith("附件：1．"))
            second = next(p for p in formatted.paragraphs if p.text.startswith("2．"))
            self.assertTrue(any(text.startswith("5．") for text in texts))
            first_index = _paragraph_index(formatted, first)
            self.assertEqual(formatted.paragraphs[first_index - 1].text, "")
            self.assertEqual(formatted.paragraphs[first_index - 2].text, "正文结束段落。")
            self.assertEqual(_indent_chars(first, "leftChars"), 200)
            self.assertEqual(_indent_chars(first, "hangingChars"), 450)
            self.assertEqual(first.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
            for prefix in ("2．", "3．", "4．", "5．"):
                item = next(p for p in formatted.paragraphs if p.text.startswith(prefix))
                self.assertEqual(_indent_chars(item, "leftChars"), 500)
                self.assertEqual(_indent_chars(item, "hangingChars"), 150)
                self.assertEqual(item.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
                self.assertEqual(_east_asia_font(item.runs[0]), "Times New Roman")
                self.assertEqual(_east_asia_font(item.runs[1]), "仿宋_GB2312")

    def test_text_input_loads_as_document(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_txt_") as tmp:
            source = Path(tmp) / "sample.txt"
            source.write_text("关于测试的通知\n各部门：\n正文。", encoding="utf-8")

            doc, kind = load_document(source)

            self.assertEqual(kind, "txt")
            self.assertEqual([p.text for p in doc.paragraphs[:3]], ["关于测试的通知", "各部门：", "正文。"])


def _indent_chars(paragraph, attr: str):
    p_pr = paragraph._p.pPr
    ind = p_pr.ind if p_pr is not None else None
    return int(ind.get(qn(f"w:{attr}"))) if ind is not None and ind.get(qn(f"w:{attr}")) else None


def _east_asia_font(run) -> str:
    r_pr = run._element.rPr
    r_fonts = r_pr.rFonts if r_pr is not None else None
    return r_fonts.get(qn("w:eastAsia"), "") if r_fonts is not None else ""


def _ascii_font(run) -> str:
    r_pr = run._element.rPr
    r_fonts = r_pr.rFonts if r_pr is not None else None
    return r_fonts.get(qn("w:ascii"), "") if r_fonts is not None else ""


def _run_spacing(run) -> int:
    r_pr = run._element.rPr
    spacing = r_pr.find(qn("w:spacing")) if r_pr is not None else None
    return int(spacing.get(qn("w:val"))) if spacing is not None else 0


def _paragraph_with_text(doc, text: str):
    return next(paragraph for paragraph in doc.paragraphs if paragraph.text.rstrip() == text)


def _paragraph_with_prefix(doc, prefix: str):
    return next(paragraph for paragraph in doc.paragraphs if paragraph.text.startswith(prefix))


def _paragraph_index(doc, target) -> int:
    return next(index for index, paragraph in enumerate(doc.paragraphs) if paragraph._p is target._p)


def _blank_count_between(doc, before, after) -> int:
    start = _paragraph_index(doc, before)
    end = _paragraph_index(doc, after)
    return sum(1 for paragraph in doc.paragraphs[start + 1 : end] if not paragraph.text and _floating_line(paragraph) is None)


def _blank_count_before(doc, target) -> int:
    index = _paragraph_index(doc, target)
    count = 0
    for paragraph in reversed(doc.paragraphs[:index]):
        if paragraph.text or _floating_line(paragraph) is not None:
            break
        count += 1
    return count


def _blank_count_after(doc, target) -> int:
    index = _paragraph_index(doc, target)
    count = 0
    for paragraph in doc.paragraphs[index + 1 :]:
        if paragraph.text or _floating_line(paragraph) is not None:
            break
        count += 1
    return count


def _floating_line(paragraph):
    return next(iter(paragraph._p.iter("{urn:schemas-microsoft-com:vml}line")), None)


def _shape_with_id(doc, shape_id: str):
    tag = "{urn:schemas-microsoft-com:vml}line"
    return next((line for line in doc._element.body.iter(tag) if line.get("id") == shape_id), None)


def _paragraph_containing_shape(doc, shape):
    return next(paragraph for paragraph in doc.paragraphs if any(line is shape for line in paragraph._p.iter(shape.tag)))


def _internal_spaces(text: str) -> int:
    match = re.search(r"办公室( +)\d{4}年", text)
    return len(match.group(1)) if match else 0


def _frame_value(paragraph, attr: str):
    p_pr = paragraph._p.pPr
    frame = p_pr.find(qn("w:framePr")) if p_pr is not None else None
    return frame.get(qn(f"w:{attr}")) if frame is not None else None


if __name__ == "__main__":
    unittest.main()
