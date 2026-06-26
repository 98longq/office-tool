"""Generate complete document fixtures for final visual verification."""

from __future__ import annotations

from pathlib import Path
import sys

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from office_tool.config import OfficeToolConfig
from office_tool.formatter import OfficialDocumentFormatter
from office_tool.services import format_document_path


OUTPUT_ROOT = ROOT / "tests" / "artifacts" / "verification"


def build(name: str, profile: str, paragraphs: list[str]) -> Path:
    source_dir = OUTPUT_ROOT / "source"
    output_dir = OUTPUT_ROOT / "docx"
    source_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)
    source = source_dir / f"{name}_source.docx"
    output = output_dir / f"{name}.docx"
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(source)
    config = OfficeToolConfig()
    config.audit.profile = profile
    OfficialDocumentFormatter(config).format_path(source, output)
    return output


def build_generated(name: str, profile: str, paragraphs: list[str], values: dict[str, object]) -> Path:
    source_dir = OUTPUT_ROOT / "source"
    output_dir = OUTPUT_ROOT / "docx"
    source_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)
    source = source_dir / f"{name}_source.docx"
    output = output_dir / f"{name}.docx"
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(source)
    config = OfficeToolConfig()
    config.audit.profile = profile
    for key, value in values.items():
        setattr(config.generation, key, value)
    result = format_document_path(source, output, config)
    if not result.ok:
        raise RuntimeError(result.error)
    return output


def main() -> None:
    standard = [
        "关于开展2026年公文质量提升专项工作的通知",
        "（2026年6月23日）",
        "所属各单位、机关各部门：",
        "为进一步提升公文办理质量，规范文件起草、校核、印发和归档流程，现就开展公文质量提升专项工作通知如下。",
        "一、总体要求",
        "各单位要坚持问题导向和结果导向，把文字质量、格式规范和办理时效贯穿公文处理全过程，形成职责清晰、衔接顺畅、运行规范的工作机制。",
        "（一）压实工作责任。各部门主要负责人要加强把关，明确专人负责公文起草和校核，重要文稿应当履行集体研究程序。",
        "（二）强化业务学习。围绕常见错情、格式标准和行文规则组织专题学习，及时纠正标题、层级、附件和落款等方面的问题。",
        "1.建立问题台账。对日常办理中发现的问题分类记录，明确整改措施、责任人员和完成时限。",
        "2.开展交叉检查。每季度组织一次公文质量互查，重点检查文字表述、数字日期、附件引用和版式设置。",
        "二、工作安排",
    ]
    standard += [
        f"第{i}项工作安排应当结合实际细化落实，形成可检查、可追溯的闭环记录，并于月底前报送阶段进展情况。"
        for i in range(1, 6)
    ]
    standard += [
        "三、有关要求",
        "各单位要严格执行保密纪律，不得通过非工作渠道传递内部文件。工作中遇到的重要情况，应当及时报告办公室协调处理。",
        "附件：1.公文质量检查项目及整改工作责任分工明细表。",
        "2.公文办理常见格式问题识别方法和具体纠正措施完整参考清单。",
        "3.年度重点任务推进情况以及阶段目标完成情况统计汇总明细表。",
        "4.各部门协同办理事项责任人员和时间节点安排工作台账。",
        "5.现场检查发现问题整改闭环以及复核确认情况报告材料。",
        "某某集团有限公司办公室",
        "2026年6月23日",
    ]

    red_head = [
        "内部资料不得外传",
        "某某集团有限公司文件",
        "某发〔2026〕68号",
        "关于进一步规范安全生产信息报送工作的通知",
        "所属各单位、机关各部门：",
        "为提升安全生产信息报送的及时性、准确性和完整性，现就有关工作要求通知如下。",
        "一、统一报送范围",
        "各单位应当全面收集生产运行、隐患排查、风险管控和应急处置等信息，确保内容真实、数据准确、口径一致。",
        "（一）明确报送责任。各单位应当确定信息报送责任部门和具体联系人，重要事项由分管负责人校核后报送。",
        "（二）规范报送内容。信息材料应当写明时间、地点、事件经过、处置措施和后续安排，不得迟报、漏报或者瞒报。",
        "二、严格报送时限",
    ]
    red_head += [
        f"第{i}类信息应当按照规定时限报送，情况发生变化时及时续报，处置结束后形成完整书面材料。"
        for i in range(1, 6)
    ]
    red_head += [
        "三、加强监督检查",
        "办公室将定期汇总报送情况，对内容不完整、格式不规范或者多次迟报的单位予以提醒，并督促限期整改。",
        "某某集团有限公司",
        "2026年6月23日",
        "抄送：上级主管单位，各所属单位。",
        "某某集团有限公司办公室    2026年6月23日印发",
    ]

    letter = [
        "某某集团有限公司",
        "某部门函〔2026〕36号",
        "关于商请协助开展设备联合检查工作的函",
        "某某市有关单位：",
        "为做好重点设备安全运行保障工作，我公司拟联合开展设备专项检查，现商请贵单位协助做好有关工作。",
        "一、检查时间",
        "计划于2026年7月上旬开展现场检查，具体时间由双方联系人提前沟通确定。",
        "二、检查内容",
    ]
    letter += [
        f"检查事项{i}主要核对设备运行记录、维护保养资料和隐患整改情况，对发现的问题现场交换意见并形成工作记录。"
        for i in range(1, 9)
    ]
    letter += [
        "三、协作事项",
        "请贵单位安排熟悉设备情况的人员参加，并提前准备相关台账资料。检查期间如有调整，双方及时协商处理。",
        "妥否，请函复。",
        "某某集团有限公司",
        "2026年6月23日",
        "（内部资料　　不得外传）",
        "（联系人：张三；电话：12345678）",
    ]

    meeting = [
        "内部资料不得外传",
        "会议纪要",
        "（18）",
        "某某集团有限公司办公室    2026年6月23日",
        "关于研究部署下半年重点工作的会议纪要",
        "2026年6月23日，公司召开专题会议，研究部署下半年重点工作。会议听取了有关部门汇报，并形成如下意见。",
        "一、提高思想认识",
        "各部门要准确把握年度目标任务，增强责任意识和协同意识，及时解决工作推进中的堵点问题。",
        "二、抓好任务落实",
    ]
    meeting += [
        f"第{i}项重点任务由责任部门细化工作清单，明确完成时限和质量要求，每月报告推进情况。"
        for i in range(1, 3)
    ]
    meeting += [
        "三、强化协同配合",
        "牵头部门要加强统筹协调，相关部门主动配合，重大事项及时提交会议研究。",
        "四、严格监督检查",
        "办公室按月跟踪重点任务完成情况，对进度滞后的事项及时提醒并督促整改。",
        "出席：张三、李四、王五、赵六、钱七、孙八、周九、吴十、郑十一、王十二、冯十三、陈十四、褚十五、卫十六",
        "分送：公司领导，各部门，各所属单位",
    ]

    outputs = [
        build("01_standard_document", "standard", standard),
        build("02_red_head_document", "red_head", red_head),
        build("03_letter_document", "letter_head", letter),
        build("04_meeting_minutes", "meeting_minutes", meeting),
        build_generated(
            "05_generated_red_head",
            "red_head",
            [
                "关于加强现场安全管理工作的通知",
                "所属各单位、机关各部门：",
                "为进一步规范现场安全管理，现就有关事项通知如下。",
                "一、落实工作责任",
                "各单位要明确责任人员，细化工作措施，确保各项要求落实到位。",
                "某某集团有限公司",
                "2026年6月23日",
            ],
            {
                "add_red_head": True,
                "add_imprint": True,
                "red_head_title": "某某集团有限公司文件",
                "document_number": "某发〔2026〕68号",
                "copy_to": "上级主管单位，各所属单位",
                "print_organization": "某某集团有限公司办公室",
                "print_date": "2026年6月23日",
            },
        ),
        build_generated(
            "06_generated_letter",
            "letter_head",
            [
                "关于商请协助开展联合检查工作的函",
                "某某市有关单位：",
                "为做好重点设备安全运行保障工作，商请贵单位协助开展联合检查。",
                "某某集团有限公司",
                "2026年6月23日",
            ],
            {
                "add_red_head": True,
                "red_head_title": "某某集团有限公司",
                "document_number": "某部门函〔2026〕42号",
            },
        ),
        build_generated(
            "07_generated_meeting_minutes",
            "meeting_minutes",
            [
                "关于研究部署下半年重点工作的会议纪要",
                "会议研究了有关事项，并形成如下意见。",
                "一、提高思想认识",
                "各部门要增强责任意识，按照会议要求抓好工作落实。",
                "出席：张三、李四、王五",
            ],
            {
                "add_red_head": True,
                "add_imprint": True,
                "meeting_number": "18",
                "meeting_organization": "某某集团有限公司办公室",
                "meeting_date": "2026年6月23日",
                "distribution": "公司领导，各部门，各所属单位",
            },
        ),
        build_generated(
            "08_generated_simple_imprint",
            "red_head",
            [
                "关于开展公文质量专项检查工作的通知",
                "所属各单位、机关各部门：",
                "请按照要求认真开展自查整改，并按时报送有关材料。",
                "某某集团有限公司",
                "2026年6月23日",
            ],
            {
                "add_imprint": True,
                "copy_to": "",
                "print_organization": "某某集团有限公司办公室",
                "print_date": "2026年6月23日",
            },
        ),
    ]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
