import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from office_tool.io import load_document
from office_tool.legacy_doc import AUTOMATION_SECURITY_FORCE_DISABLE, DOCX_FILE_FORMAT, convert_legacy_doc


class _FakePythonCom:
    def __init__(self):
        self.initialized = 0
        self.uninitialized = 0

    def CoInitialize(self):
        self.initialized += 1

    def CoUninitialize(self):
        self.uninitialized += 1


class _FakeDocument:
    def __init__(self):
        self.closed = False
        self.saved_format = None

    def SaveAs2(self, destination, file_format):
        self.saved_format = file_format
        Path(destination).write_bytes(b"converted-docx")

    def Close(self, _save_changes):
        self.closed = True


class _FakeDocuments:
    def __init__(self, document):
        self.document = document
        self.open_args = None

    def Open(self, *args):
        self.open_args = args
        return self.document


class _FakeApplication:
    def __init__(self):
        self.document = _FakeDocument()
        self.Documents = _FakeDocuments(self.document)
        self.quit_called = False
        self.Visible = True
        self.DisplayAlerts = 1
        self.AutomationSecurity = 0

    def Quit(self):
        self.quit_called = True


class _FakeClient:
    def __init__(self, application):
        self.application = application
        self.attempts = []

    def DispatchEx(self, prog_id):
        self.attempts.append(prog_id)
        if prog_id == "Word.Application":
            raise RuntimeError("Word unavailable")
        return self.application


class LegacyDocTests(unittest.TestCase):
    def test_converter_falls_back_to_wps_and_disables_macros(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_legacy_doc_test_") as tmp:
            root = Path(tmp)
            source = root / "sample.doc"
            destination = root / "sample.docx"
            source.write_bytes(b"legacy-doc")
            pythoncom = _FakePythonCom()
            application = _FakeApplication()
            client = _FakeClient(application)

            with patch("office_tool.legacy_doc.sys.platform", "win32"), patch(
                "office_tool.legacy_doc._load_com_modules",
                return_value=(pythoncom, client),
            ):
                result = convert_legacy_doc(source, destination)

            self.assertEqual(result, destination.resolve())
            self.assertEqual(client.attempts[:2], ["Word.Application", "kwps.Application"])
            self.assertEqual(application.AutomationSecurity, AUTOMATION_SECURITY_FORCE_DISABLE)
            self.assertEqual(application.document.saved_format, DOCX_FILE_FORMAT)
            self.assertTrue(application.document.closed)
            self.assertTrue(application.quit_called)
            self.assertEqual(pythoncom.initialized, 1)
            self.assertEqual(pythoncom.uninitialized, 1)

    def test_conversion_failure_uses_chinese_stage_messages(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_legacy_error_test_") as tmp:
            root = Path(tmp)
            source = root / "sample.doc"
            destination = root / "sample.docx"
            source.write_bytes(b"legacy-doc")
            pythoncom = _FakePythonCom()
            client = _FakeClient(_FakeApplication())

            with patch("office_tool.legacy_doc.sys.platform", "win32"), patch(
                "office_tool.legacy_doc._load_com_modules",
                return_value=(pythoncom, client),
            ), patch("office_tool.legacy_doc.COM_BACKENDS", (("Microsoft Word", "Word.Application"),)):
                with self.assertRaisesRegex(Exception, "无法转换 .doc 文件.*启动程序失败"):
                    convert_legacy_doc(source, destination)

    def test_load_document_converts_doc_in_temporary_directory(self):
        with tempfile.TemporaryDirectory(prefix="office_tool_legacy_load_test_") as tmp:
            source = Path(tmp) / "sample.doc"
            source.write_bytes(b"legacy-doc")

            def fake_convert(_source, destination):
                converted = Document()
                converted.add_paragraph("旧版文档内容")
                converted.save(destination)
                return Path(destination)

            with patch("office_tool.io.convert_legacy_doc", side_effect=fake_convert) as converter:
                document, kind = load_document(source)

            self.assertEqual(kind, "doc")
            self.assertEqual(document.paragraphs[0].text, "旧版文档内容")
            converter.assert_called_once()


if __name__ == "__main__":
    unittest.main()
