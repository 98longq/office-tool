"""Input loading helpers."""

from __future__ import annotations

import re
import tempfile
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument

from .legacy_doc import convert_legacy_doc

SUPPORTED_INPUTS = {".doc", ".docx", ".txt", ".md"}


class UnsupportedInputError(ValueError):
    pass


def load_document(path: str | Path) -> tuple[DocxDocument, str]:
    source = Path(path).expanduser().resolve()
    suffix = source.suffix.lower()
    if suffix == ".doc":
        with tempfile.TemporaryDirectory(prefix="office_tool_legacy_doc_") as tmp:
            converted = Path(tmp) / f"{source.stem}.docx"
            convert_legacy_doc(source, converted)
            return Document(converted), "doc"
    if suffix == ".docx":
        return Document(source), "docx"
    if suffix == ".txt":
        return _document_from_text(source.read_text(encoding="utf-8")), "txt"
    if suffix == ".md":
        return _document_from_text(_clean_markdown(source.read_text(encoding="utf-8"))), "md"
    raise UnsupportedInputError(
        f"当前版本支持 .doc/.docx/.txt/.md，暂不处理 {suffix or '无扩展名'}。"
    )


def _document_from_text(text: str) -> DocxDocument:
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    if not doc.paragraphs:
        doc.add_paragraph("")
    return doc


def _clean_markdown(text: str) -> str:
    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if re.match(r"^\s*[-*_]{3,}\s*$", stripped):
            continue
        line = re.sub(r"^\s{0,3}#{1,6}\s+", "", line)
        line = re.sub(r"^\s*>\s?", "", line)
        line = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", line)
        line = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line)
        line = re.sub(r"(\*\*|__)(.*?)\1", r"\2", line)
        line = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", line)
        line = re.sub(r"`([^`]+)`", r"\1", line)
        lines.append(line)
    return "\n".join(lines)
