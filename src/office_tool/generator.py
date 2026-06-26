"""Conservative optional content generation for red-head documents."""

from __future__ import annotations

import re
from dataclasses import dataclass, field

from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from .audit import OfficialDocumentAuditor
from .config import OfficeToolConfig


@dataclass
class GenerationResult:
    generated: list[str] = field(default_factory=list)
    skipped: list[str] = field(default_factory=list)

    def to_stats(self) -> dict[str, object]:
        return {
            "generated_content": list(self.generated),
            "skipped_generation": list(self.skipped),
        }


class OfficialDocumentContentGenerator:
    """Generate only wholly absent structures whose values were supplied by the user."""

    RED_HEAD_PROFILES = {"red_head", "letter_head", "meeting_minutes"}

    def __init__(self, config: OfficeToolConfig):
        self.config = config
        self.auditor = OfficialDocumentAuditor(config)

    def apply(self, doc: DocxDocument) -> GenerationResult:
        result = GenerationResult()
        options = self.config.generation
        profile = self.config.audit.profile
        if profile not in self.RED_HEAD_PROFILES:
            return result

        structure = self.auditor.detect_structure(self.auditor._paragraph_items(doc.paragraphs))
        if options.add_red_head:
            self._add_red_head(doc, profile, structure, result)

        # Re-detect because newly generated paragraphs change all indices.
        structure = self.auditor.detect_structure(self.auditor._paragraph_items(doc.paragraphs))
        if options.add_imprint and profile != "letter_head":
            self._add_imprint(doc, profile, structure, result)
        return result

    def _add_red_head(self, doc, profile, structure, result: GenerationResult) -> None:
        options = self.config.generation
        if profile == "meeting_minutes":
            existing = [
                structure.red_head,
                structure.meeting_number,
                structure.meeting_issue_line,
            ]
            if any(index is not None for index in existing):
                result.skipped.append("red_head_existing_or_partial")
                return
            self._require(
                ("会议期号", options.meeting_number),
                ("编发单位", options.meeting_organization),
                ("编发日期", options.meeting_date),
            )
            paragraphs = [
                "内部资料不得外传",
                "会议纪要",
                self._meeting_number(options.meeting_number),
                f"{options.meeting_organization}    {options.meeting_date}",
            ]
            self._prepend_paragraphs(doc, paragraphs)
            result.generated.append("red_head")
            return

        if profile == "letter_head":
            self._add_letter_red_head(doc, structure, result)
            return

        self._add_standard_red_head(doc, structure, result)

    def _add_standard_red_head(self, doc: DocxDocument, structure, result: GenerationResult) -> None:
        options = self.config.generation
        self._require_missing_red_head_values(structure, options)
        changed = False
        if structure.internal_notice is None:
            self._insert_paragraph(doc, 0, "内部资料不得外传")
            changed = True
            structure = self.auditor.detect_structure(self.auditor._paragraph_items(doc.paragraphs))
        if structure.red_head is None:
            anchor = structure.document_number if structure.document_number is not None else 1
            self._insert_paragraph(doc, max(0, anchor), options.red_head_title.strip())
            changed = True
            structure = self.auditor.detect_structure(self.auditor._paragraph_items(doc.paragraphs))
        if structure.document_number is None:
            anchor = (structure.red_head + 1) if structure.red_head is not None else len(doc.paragraphs)
            self._insert_paragraph(doc, anchor, options.document_number.strip())
            changed = True
        if changed:
            result.generated.append("red_head")
        else:
            result.skipped.append("red_head_existing")

    def _add_letter_red_head(self, doc: DocxDocument, structure, result: GenerationResult) -> None:
        options = self.config.generation
        self._require_missing_red_head_values(structure, options)
        changed = False
        if structure.red_head is None:
            anchor = structure.document_number if structure.document_number is not None else 0
            self._insert_paragraph(doc, max(0, anchor), options.red_head_title.strip())
            changed = True
            structure = self.auditor.detect_structure(self.auditor._paragraph_items(doc.paragraphs))
        if structure.document_number is None:
            anchor = (structure.red_head + 1) if structure.red_head is not None else 1
            self._insert_paragraph(doc, anchor, options.document_number.strip())
            changed = True
        self._append_letter_notes(doc)
        if changed:
            result.generated.append("red_head")
        else:
            result.skipped.append("red_head_existing")

    def _add_imprint(self, doc, profile, structure, result: GenerationResult) -> None:
        if profile == "meeting_minutes":
            if structure.distribution is not None:
                result.skipped.append("distribution_existing")
                return
            value = re.sub(r"^分送\s*[:：]\s*", "", self.config.generation.distribution.strip())
            self._require(("分送内容", value))
            doc.add_paragraph(f"分送：{value}")
            result.generated.append("distribution")
            return

        if structure.copy_to is not None or structure.print_org_date is not None or structure.simple_imprint is not None:
            result.skipped.append("imprint_existing_or_partial")
            return
        options = self.config.generation
        copy_to = re.sub(r"^(?:抄送|抄报|发送)\s*[:：]\s*", "", options.copy_to.strip())
        print_date = re.sub(r"印发\s*$", "", options.print_date.strip())
        self._require(
            ("印发单位", options.print_organization),
            ("印发日期", print_date),
        )
        if copy_to:
            doc.add_paragraph(f"抄送：{copy_to}")
            doc.add_paragraph(f"{options.print_organization.strip()}    {print_date}印发")
            result.generated.append("imprint")
        else:
            paragraph = doc.add_paragraph(f"{options.print_organization.strip()}    {print_date}")
            self._set_paragraph_style_id(paragraph, "OfficeToolGeneratedSimpleImprint")
            result.generated.append("simple_imprint")

    def _append_letter_notes(self, doc: DocxDocument) -> None:
        texts = [paragraph.text.strip() for paragraph in doc.paragraphs]
        if not any("内部资料" in text and "不得外传" in text for text in texts):
            doc.add_paragraph("（内部资料　　不得外传）")

    @staticmethod
    def _set_paragraph_style_id(paragraph: Paragraph, style_id: str) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        p_style = p_pr.get_or_add_pStyle()
        p_style.set(qn("w:val"), style_id)

    @staticmethod
    def _meeting_number(value: str) -> str:
        stripped = value.strip().strip("（）()")
        return f"（{stripped}）"

    @staticmethod
    def _require(*values: tuple[str, str]) -> None:
        missing = [label for label, value in values if not str(value).strip()]
        if missing:
            raise ValueError("添加内容前请填写：" + "、".join(missing))

    @classmethod
    def _require_missing_red_head_values(cls, structure, options) -> None:
        required: list[tuple[str, str]] = []
        if structure.red_head is None:
            required.append(("红头名称", options.red_head_title))
        if structure.document_number is None:
            required.append(("发文字号", options.document_number))
        if required:
            cls._require(*required)

    @staticmethod
    def _prepend_paragraphs(doc: DocxDocument, texts: list[str]) -> None:
        body = doc._element.body
        for index, text in enumerate(texts):
            element = OxmlElement("w:p")
            body.insert(index, element)
            Paragraph(element, doc._body).add_run(text)

    @staticmethod
    def _insert_paragraph(doc: DocxDocument, index: int, text: str) -> Paragraph:
        body = doc._element.body
        element = OxmlElement("w:p")
        body.insert(max(0, min(index, len(doc.paragraphs))), element)
        paragraph = Paragraph(element, doc._body)
        paragraph.add_run(text)
        return paragraph
