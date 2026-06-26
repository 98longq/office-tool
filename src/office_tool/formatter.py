"""Official document formatting engine."""

from __future__ import annotations

import re
import unicodedata
from pathlib import Path

from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Twips
from docx.table import Table
from docx.text.paragraph import Paragraph

from .audit import OfficialDocumentAuditor
from .config import OfficeToolConfig
from .docx_utils import (
    add_floating_line,
    add_page_number,
    apply_font,
    apply_style_to_paragraph,
    clear_paragraph_frame,
    set_side_indent_chars,
    set_document_grid,
    set_even_and_odd_headers,
    add_body_text_box,
)
from .io import load_document
from .generator import OfficialDocumentContentGenerator
from .models import AuditReport
from .patterns import (
    RE_ATTACHMENT_NOTE,
    RE_DISTRIBUTION,
    RE_MEETING_ATTENDEES,
    RE_MEETING_ISSUE_LINE,
    RE_MEETING_NUMBER,
    RE_MEETING_RED_HEAD,
    RE_REGULATION_ARTICLE_PARTS,
    RE_REGULATION_CHAPTER,
    RE_TITLE_DATE_LINE,
    heading_role,
    is_heading,
)


class OfficialDocumentFormatter:
    """Audit and format documents using configurable official-document rules."""

    def __init__(self, config: OfficeToolConfig | None = None):
        self.config = config or OfficeToolConfig()
        self.auditor = OfficialDocumentAuditor(self.config)

    def format_path(self, input_path: str | Path, output_path: str | Path) -> AuditReport:
        doc, _kind = load_document(input_path)
        report = self.format_document(doc)
        output = Path(output_path).expanduser().resolve()
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output)
        return report

    def format_document(self, doc: DocxDocument) -> AuditReport:
        generation = OfficialDocumentContentGenerator(self.config).apply(doc)
        report = self.auditor.audit_document(doc)
        if self.config.format.apply_page_setup:
            self._apply_page_setup(doc, report)
        if self.config.format.apply_styles:
            self._apply_styles(doc, report)
        if self.config.format.add_page_number:
            self._apply_page_number(doc, report)
        if report.is_letter_head:
            self._apply_letter_header(doc, report)
        final_report = self.auditor.audit_document(doc)
        final_report.stats.update(generation.to_stats())
        return final_report

    def _apply_page_setup(self, doc: DocxDocument, report: AuditReport) -> None:
        page = self.config.page
        body_style = self.config.styles["body"]
        self._apply_normal_style(doc, body_style.font, body_style.size_pt, body_style.latin_font)
        for section in doc.sections:
            section.page_width = Cm(page.paper_width_cm)
            section.page_height = Cm(page.paper_height_cm)
            section.top_margin = Cm(page.margin_top_cm)
            section.bottom_margin = Cm(2.5 if report.is_letter_head else page.margin_bottom_cm)
            section.left_margin = Cm(page.margin_left_cm)
            section.right_margin = Cm(page.margin_right_cm)
            section.footer_distance = Cm(page.footer_distance_cm)
            if self.config.format.apply_document_grid:
                self._fit_grid_width(section, page.chars_per_line, page.grid_char_space_pt)
                grid_char_space_pt = self._effective_grid_char_space(
                    section,
                    page.chars_per_line,
                    page.grid_char_space_pt,
                )
                set_document_grid(
                    section,
                    page.chars_per_line,
                    page.lines_per_page,
                    grid_char_space_pt,
                    page.grid_line_pitch_pt,
                    body_style.size_pt,
                )

    @staticmethod
    def _fit_grid_width(section, chars_per_line: int, char_space_pt: float) -> None:
        required_width = int(round(chars_per_line * char_space_pt * 20))
        available_width = section.page_width.twips - section.left_margin.twips - section.right_margin.twips
        deficit = required_width - available_width
        if deficit <= 0:
            return

        # Word stores dimensions in integer twips. Split the sub-0.01 cm
        # compensation across both margins so its UI still shows 2.80/2.60 cm.
        left_adjustment = deficit // 2
        right_adjustment = deficit - left_adjustment
        section.left_margin = Twips(section.left_margin.twips - left_adjustment)
        section.right_margin = Twips(section.right_margin.twips - right_adjustment)

    @staticmethod
    def _effective_grid_char_space(section, chars_per_line: int, requested_pt: float) -> float:
        available_width = section.page_width.twips - section.left_margin.twips - section.right_margin.twips
        # Word rounds a grid that exactly fills the text area down by one
        # character. Reserve two twips while retaining a displayed value of
        # 15.8 pt in its one-decimal UI field.
        max_pitch_pt = (available_width - 2) / (chars_per_line * 20)
        return min(requested_pt, max_pitch_pt)

    def _apply_styles(self, doc: DocxDocument, report: AuditReport) -> None:
        role_by_index = self._roles_from_report(report)
        paragraph_by_index = dict(enumerate(doc.paragraphs))
        date_index = self._first_index_for(role_by_index, "date")
        if date_index is not None:
            doc.paragraphs[date_index].text = self._normalize_document_date_text(doc.paragraphs[date_index].text)
        date_text = doc.paragraphs[date_index].text.strip() if date_index is not None else ""
        title_index = self._first_index_for(role_by_index, "title")
        title_block = self._title_block_indices(doc.paragraphs, title_index, role_by_index)

        for index, paragraph in enumerate(doc.paragraphs):
            if not paragraph.text.strip():
                continue
            role = self._role_for_paragraph(index, paragraph.text, report.profile, role_by_index, title_block)
            if role == "internal_notice":
                if report.is_letter_head:
                    self._normalize_letter_internal_notice(paragraph)
                else:
                    self._normalize_internal_notice(paragraph)
            elif role == "red_head" and report.is_meeting_minutes:
                paragraph.text = "会 议 纪 要"
            elif role == "meeting_number":
                self._normalize_meeting_number(paragraph)
            elif role == "meeting_issue_line":
                self._normalize_meeting_issue_line(paragraph)
            elif role == "meeting_attendees":
                self._normalize_paragraph_text(paragraph)
            elif role == "distribution":
                self._normalize_distribution(paragraph)
            elif role == "letter_contact":
                self._normalize_letter_contact(paragraph)
            elif role == "regulation_chapter":
                self._normalize_regulation_chapter(paragraph)
            elif role == "regulation_article":
                self._normalize_regulation_article(paragraph)
            elif role == "copy_to":
                self._normalize_copy_to_text(paragraph)
            elif role == "print_org_date":
                self._normalize_print_org_date_text(paragraph)
            elif role == "simple_imprint":
                self._normalize_simple_imprint_text(paragraph)
            else:
                self._normalize_paragraph_text(paragraph)
            if role in {"attachment_note", "attachment_item"}:
                self._normalize_attachment_text(paragraph)
            if role in {"h2", "h3"}:
                self._ensure_heading_period(paragraph)
            if report.is_letter_head and role == "internal_notice":
                style_name = "letter_internal_notice"
            elif report.is_meeting_minutes and role == "red_head":
                style_name = "meeting_red_head"
            else:
                style_name = self._style_for_role(role)
            style = self.config.styles[style_name]
            if role == "h2" and self._apply_mixed_h2_style(paragraph):
                continue
            if role == "regulation_article" and self._apply_regulation_article_style(paragraph):
                continue
            if role == "meeting_attendees" and self._apply_meeting_attendees_style(paragraph):
                continue
            apply_style_to_paragraph(
                paragraph,
                style,
                preserve_bold_italic=self.config.format.preserve_existing_bold_italic,
            )
            if role in {"attachment_note", "attachment_item"}:
                self._apply_attachment_hanging_layout(paragraph, role)
            if role == "date":
                self._apply_date_layout(paragraph)
            elif role == "signatory":
                self._apply_signatory_layout(paragraph, date_text)
            elif role == "red_head":
                if not report.is_letter_head and not report.is_meeting_minutes:
                    self._fit_red_head_text(paragraph)
            if role != "signatory":
                self._force_ascii_runs_to_latin_font(paragraph, style)
        self._normalize_blank_paragraphs(doc, role_by_index, title_block, paragraph_by_index)
        if report.is_red_head and not report.is_letter_head:
            if report.is_meeting_minutes:
                self._apply_meeting_minutes_layout(role_by_index, paragraph_by_index)
            else:
                self._apply_red_head_layout(role_by_index, paragraph_by_index)
                self._apply_regulation_layout(role_by_index, paragraph_by_index)
        if report.is_letter_head:
            self._apply_letter_head_layout(role_by_index, paragraph_by_index)
        if self.config.format.draw_imprint_lines:
            self._apply_imprint_layout(doc, role_by_index, paragraph_by_index)
            self._apply_simple_imprint_layout(doc, role_by_index, paragraph_by_index)
        self._apply_distribution_layout(doc, role_by_index, paragraph_by_index)
        self._normalize_all_line_spacing(doc)

    def _normalize_blank_paragraphs(
        self,
        doc: DocxDocument,
        role_by_index: dict[int, str],
        title_block: set[int] | None = None,
        paragraph_by_index: dict[int, Paragraph] | None = None,
    ) -> None:
        paragraph_by_index = paragraph_by_index or dict(enumerate(doc.paragraphs))
        title_index = self._first_index_for(role_by_index, "title")
        signatory_index = self._first_index_for(role_by_index, "signatory")
        attachment_index = self._first_index_for(role_by_index, "attachment_note")
        title_anchor_index = max(title_block) if title_block else title_index
        title = (
            paragraph_by_index.get(title_anchor_index)
            if title_anchor_index is not None
            else None
        )
        signatory = (
            paragraph_by_index.get(signatory_index)
            if signatory_index is not None
            else None
        )
        attachment = (
            paragraph_by_index.get(attachment_index)
            if attachment_index is not None
            else None
        )
        if title is not None:
            self._set_blank_after(title, 1)
        if signatory is not None:
            self._set_blank_before(signatory, 2)
        if attachment is not None:
            self._set_blank_before(attachment, 1)

    @staticmethod
    def _is_blank_paragraph_element(p_element) -> bool:
        return not "".join(node.text or "" for node in p_element.iter(qn("w:t"))).strip()

    @classmethod
    def _set_blank_after(cls, paragraph: Paragraph, count: int) -> None:
        current = paragraph._p.getnext()
        while current is not None and current.tag == qn("w:p") and cls._is_blank_paragraph_element(current):
            next_element = current.getnext()
            current.getparent().remove(current)
            current = next_element
        anchor = paragraph._p
        for _ in range(count):
            blank = OxmlElement("w:p")
            anchor.addnext(blank)
            anchor = blank

    @classmethod
    def _set_blank_before(cls, paragraph: Paragraph, count: int) -> None:
        current = paragraph._p.getprevious()
        while current is not None and current.tag == qn("w:p") and cls._is_blank_paragraph_element(current):
            previous_element = current.getprevious()
            current.getparent().remove(current)
            current = previous_element
        for _ in range(count):
            blank = OxmlElement("w:p")
            paragraph._p.addprevious(blank)

    @staticmethod
    def _text_width_chars(text: str) -> float:
        width = 0.0
        for char in text.strip():
            if char.isspace():
                continue
            width += 1.0 if unicodedata.east_asian_width(char) in {"F", "W"} else 0.5
        return width

    def _title_block_indices(
        self, paragraphs: list[Paragraph], title_index: int | None, role_by_index: dict[int, str]
    ) -> set[int]:
        if title_index is None:
            return set()
        indices = {title_index}
        for index in range(title_index - 1, max(-1, title_index - 3), -1):
            if role_by_index.get(index) not in {None, "title"}:
                break
            text = paragraphs[index].text.strip()
            if self._looks_like_title_continuation(text):
                indices.add(index)
                continue
            break
        for index in range(title_index + 1, min(len(paragraphs), title_index + 4)):
            text = paragraphs[index].text.strip()
            if not text:
                continue
            if RE_TITLE_DATE_LINE.match(text) or self._looks_like_title_continuation(text):
                indices.add(index)
                continue
            break
        return indices

    @staticmethod
    def _looks_like_title_continuation(text: str) -> bool:
        if not text or len(text) > 40:
            return False
        if RE_TITLE_DATE_LINE.match(text) or is_heading(text) or RE_ATTACHMENT_NOTE.match(text):
            return False
        if text.endswith(("：", ":", "。", "；", ";", "！", "？")):
            return False
        return True

    def _role_for_paragraph(
        self,
        index: int,
        text: str,
        profile: str,
        role_by_index: dict[int, str],
        title_block: set[int],
    ) -> str:
        if index in title_block:
            return "title_date" if RE_TITLE_DATE_LINE.match(text.strip()) else "title"
        reported_role = role_by_index.get(index)
        structural_roles = {
            "internal_notice", "red_head", "document_number", "signer", "main_send",
            "attachment_note", "signatory", "date", "copy_to", "print_org_date", "simple_imprint",
            "regulation_code", "regulation_title", "regulation_chapter", "regulation_article",
            "letter_contact", "meeting_number", "meeting_issue_line", "meeting_attendees",
            "distribution",
        }
        role = (
            reported_role
            if reported_role in structural_roles
            else self._derived_role(text, profile) or reported_role or "body"
        )
        if self._is_attachment_item(index, text, role_by_index):
            return "attachment_item"
        return role

    @staticmethod
    def _is_attachment_item(index: int, text: str, role_by_index: dict[int, str]) -> bool:
        if not re.match(r"^\s*\d{1,2}[．.、]", text):
            return False
        previous_attachment = [i for i, role in role_by_index.items() if role == "attachment_note" and i < index]
        previous_stop = [
            i
            for i, role in role_by_index.items()
            if role in {"signatory", "date", "copy_to", "print_org_date", "simple_imprint"} and i < index
        ]
        return bool(previous_attachment) and (not previous_stop or max(previous_attachment) > max(previous_stop))

    def _apply_mixed_h2_style(self, paragraph: Paragraph) -> bool:
        text = paragraph.text
        match = re.match(r"^([（(]\s*[一二三四五六七八九十]+\s*[）)].*?[。！？；;])(.+)$", text)
        if not match:
            return False
        heading_text, body_text = match.groups()
        if not body_text.strip():
            return False
        body_style = self.config.styles["body"]
        heading_style = self.config.styles["h2"]
        apply_style_to_paragraph(
            paragraph,
            body_style,
            preserve_bold_italic=self.config.format.preserve_existing_bold_italic,
        )
        paragraph.clear()
        heading_run = paragraph.add_run(heading_text)
        apply_font(heading_run, heading_style.font, heading_style.size_pt, heading_style.color, latin_font=heading_style.latin_font)
        heading_run.bold = heading_style.bold
        body_run = paragraph.add_run(body_text)
        apply_font(body_run, body_style.font, body_style.size_pt, body_style.color, latin_font=body_style.latin_font)
        body_run.bold = body_style.bold
        return True

    @staticmethod
    def _normalize_attachment_text(paragraph: Paragraph) -> None:
        text = paragraph.text.strip()
        text = re.sub(r"^(附件：\s*\d{1,2})[.．、]", r"\1．", text)
        text = re.sub(r"^(\d{1,2})[.．、]", r"\1．", text)
        text = re.sub(r"[。；;]\s*$", "", text)
        paragraph.text = text

    @staticmethod
    def _apply_attachment_hanging_layout(paragraph: Paragraph, role: str) -> None:
        if role == "attachment_note":
            left_chars = 2.0
            hanging_chars = 4.5
        else:
            left_chars = 5.0
            hanging_chars = 1.5
        size_pt = 16
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Pt(left_chars * size_pt)
        paragraph.paragraph_format.first_line_indent = Pt(-hanging_chars * size_pt)
        p_pr = paragraph._p.get_or_add_pPr()
        ind = p_pr.get_or_add_ind()
        ind.set(qn("w:leftChars"), str(int(left_chars * 100)))
        ind.set(qn("w:hangingChars"), str(int(hanging_chars * 100)))
        ind.attrib.pop(qn("w:firstLineChars"), None)

    @staticmethod
    def _normalize_internal_notice(paragraph: Paragraph) -> None:
        paragraph.text = "内部资料\n不得外传"

    @staticmethod
    def _normalize_letter_internal_notice(paragraph: Paragraph) -> None:
        """Letter style: (内部资料　　不得外传) — in parens with full-width spaces."""
        paragraph.text = "（内部资料　　不得外传）"

    @staticmethod
    def _normalize_letter_contact(paragraph: Paragraph) -> None:
        """Normalize contact info line for letter format."""
        text = paragraph.text.strip()
        if text.startswith("（") and text.endswith("）"):
            text = text[1:-1]
        # Normalize colon type
        text = re.sub(r"[:：]", "：", text)
        paragraph.text = f"（{text}）"

    @staticmethod
    def _normalize_meeting_number(paragraph: Paragraph) -> None:
        match = re.search(r"\d+", paragraph.text)
        if match:
            paragraph.text = f"（ {match.group()} ）"

    def _normalize_meeting_issue_line(self, paragraph: Paragraph) -> None:
        text = paragraph.text.strip()
        match = re.match(r"^(.*?)(\d{4}年\d{1,2}月\d{1,2}日)$", re.sub(r"\s+", "", text))
        if not match:
            paragraph.text = normalize_official_text(text)
            return
        organization, date_text = match.groups()
        size_pt = self.config.styles["meeting_issue_line"].size_pt
        available_pt = 15.6 / 2.54 * 72 - 2 * size_pt
        occupied_pt = (self._text_width_chars(organization) + self._text_width_chars(date_text)) * size_pt
        # A half-width space in the target fonts is close to half an em. Keep
        # a small reserve so long organization names never push the date down.
        spaces = max(2, int((available_pt - occupied_pt) // (size_pt / 2)) - 2)
        paragraph.text = organization + (" " * spaces) + date_text

    @staticmethod
    def _normalize_distribution(paragraph: Paragraph) -> None:
        text = normalize_official_text(paragraph.text.strip()).replace("、", "，")
        text = re.sub(r"^分送\s*[:：]\s*", "分送：", text)
        text = re.sub(r"[，；;\s]+$", "", text)
        paragraph.text = text if text.endswith("。") else text + "。"

    @staticmethod
    def _normalize_regulation_chapter(paragraph: Paragraph) -> None:
        match = RE_REGULATION_CHAPTER.match(paragraph.text.strip())
        if match:
            paragraph.text = f"{match.group(1)}　{match.group(2).strip()}"

    @staticmethod
    def _normalize_regulation_article(paragraph: Paragraph) -> None:
        match = RE_REGULATION_ARTICLE_PARTS.match(paragraph.text.strip())
        if match:
            paragraph.text = f"{match.group(1)}　{match.group(2).strip()}"

    @staticmethod
    def _normalize_copy_to_text(paragraph: Paragraph) -> None:
        text = normalize_official_text(paragraph.text.strip()).replace("、", "，")
        text = re.sub(r"[，。；;\s]+$", "", text)
        paragraph.text = text + "。"

    def _normalize_print_org_date_text(self, paragraph: Paragraph) -> None:
        text = re.sub(r"\s+", "", paragraph.text.strip())
        match = re.match(r"^(.*?)(\d{4}年\d{1,2}月\d{1,2}日印发)$", text)
        if not match:
            paragraph.text = normalize_official_text(paragraph.text.strip())
            return
        organization, print_date = match.groups()
        size_pt = self.config.styles["copy_to"].size_pt
        available_pt = 15.6 / 2.54 * 72 - 2 * size_pt
        text_width_pt = (self._text_width_chars(organization) + self._text_width_chars(print_date)) * size_pt
        space_width_pt = size_pt * 0.5
        spaces = max(4, int((available_pt - text_width_pt) // space_width_pt) - 2)
        paragraph.text = organization + (" " * spaces) + print_date

    def _normalize_simple_imprint_text(self, paragraph: Paragraph) -> None:
        text = re.sub(r"\s+", "", paragraph.text.strip())
        match = re.match(r"^(.*?)(\d{4}年\d{1,2}月\d{1,2}日)$", text)
        if not match:
            paragraph.text = normalize_official_text(paragraph.text.strip())
            return
        organization, date_text = match.groups()
        size_pt = self.config.styles["copy_to"].size_pt
        available_pt = 15.6 / 2.54 * 72 - 2 * size_pt
        occupied_pt = (self._text_width_chars(organization) + self._text_width_chars(date_text)) * size_pt
        spaces = max(4, int((available_pt - occupied_pt) // (size_pt / 2)) - 2)
        paragraph.text = organization + (" " * spaces) + date_text

    def _apply_regulation_article_style(self, paragraph: Paragraph) -> bool:
        match = RE_REGULATION_ARTICLE_PARTS.match(paragraph.text.strip())
        if not match:
            return False
        prefix, body = match.groups()
        paragraph_style = self.config.styles["regulation_article"]
        apply_style_to_paragraph(
            paragraph,
            paragraph_style,
            preserve_bold_italic=self.config.format.preserve_existing_bold_italic,
        )
        paragraph.clear()
        prefix_run = paragraph.add_run(prefix + "　")
        apply_font(prefix_run, "黑体", 16, latin_font="Times New Roman")
        prefix_run.bold = False
        if body:
            body_run = paragraph.add_run(body)
            apply_font(body_run, paragraph_style.font, paragraph_style.size_pt, latin_font=paragraph_style.latin_font)
            body_run.bold = False
        return True

    def _apply_meeting_attendees_style(self, paragraph: Paragraph) -> bool:
        match = re.match(r"^\s*(出席\s*[:：])(.*)$", paragraph.text, re.S)
        if not match:
            return False
        prefix, body = match.groups()
        style = self.config.styles["meeting_attendees"]
        apply_style_to_paragraph(paragraph, style, preserve_bold_italic=False)
        paragraph.clear()
        prefix_run = paragraph.add_run("出席：")
        apply_font(prefix_run, "黑体", 16, latin_font="Times New Roman")
        prefix_run.bold = False
        if body:
            body_run = paragraph.add_run(body.strip())
            apply_font(body_run, style.font, style.size_pt, latin_font=style.latin_font)
            body_run.bold = False
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Pt(2 * style.size_pt)
        paragraph.paragraph_format.first_line_indent = Pt(-3.1 * style.size_pt)
        p_pr = paragraph._p.get_or_add_pPr()
        ind = p_pr.get_or_add_ind()
        ind.set(qn("w:leftChars"), "200")
        ind.set(qn("w:hangingChars"), "310")
        ind.attrib.pop(qn("w:firstLineChars"), None)
        return True

    def _fit_red_head_text(self, paragraph: Paragraph) -> None:
        style = self.config.styles["red_head"]
        text = paragraph.text.strip().replace("\n", "")
        if text != paragraph.text:
            paragraph.text = text
            apply_style_to_paragraph(paragraph, style, preserve_bold_italic=False)

        units = max(1.0, self._text_width_chars(text))
        gaps = max(1, len(text) - 1)
        available_pt = 15.6 / 2.54 * 72 - 4
        font_size = min(48.0, style.size_pt or 42.0)
        spacing_pt = 0.0
        estimated_width = units * font_size
        if estimated_width > available_pt:
            needed_spacing = (available_pt - estimated_width) / gaps
            if needed_spacing >= -1.0:
                spacing_pt = needed_spacing
            else:
                spacing_pt = -0.8
                font_size = max(22.0, int(((available_pt - spacing_pt * gaps) / units) * 2) / 2)

        for run in paragraph.runs:
            apply_font(run, style.font, font_size, style.color, latin_font=style.latin_font)
            run.bold = False
            self._set_run_spacing(run, spacing_pt)

    @staticmethod
    def _set_run_spacing(run, spacing_pt: float) -> None:
        r_pr = run._element.get_or_add_rPr()
        spacing = r_pr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            r_pr.append(spacing)
        spacing.set(qn("w:val"), str(int(round(spacing_pt * 20))))

    def _apply_date_layout(self, paragraph: Paragraph) -> None:
        paragraph.text = self._normalize_document_date_text(paragraph.text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_side_indent_chars(paragraph, right_chars=4, size_pt=self.config.styles["date"].size_pt)

    def _apply_signatory_layout(self, paragraph: Paragraph, date_text: str) -> None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        base_text = paragraph.text.rstrip()
        paragraph.text = base_text
        apply_style_to_paragraph(
            paragraph,
            self.config.styles["signatory"],
            preserve_bold_italic=self.config.format.preserve_existing_bold_italic,
        )
        signatory_chars = self._text_width_chars(base_text)
        target_chars = 14.0
        missing_chars = max(0.0, target_chars - signatory_chars)
        trailing_spaces = int(round(missing_chars * 2))
        if trailing_spaces:
            spacer_run = paragraph.add_run(" " * max(0, trailing_spaces - 1) + "\u00a0")
            style = self.config.styles["signatory"]
            apply_font(spacer_run, style.font, style.size_pt, style.color, latin_font=style.latin_font)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_side_indent_chars(paragraph, size_pt=self.config.styles["signatory"].size_pt)

    @staticmethod
    def _normalize_document_date_text(text: str) -> str:
        stripped = text.strip()
        match = re.fullmatch(r"(\d{4})年\s*(\d{1,2})月\s*(\d{1,2})日", stripped)
        if not match:
            return stripped
        year, month, day = match.groups()
        return f"{year}年{int(month):02d}月{int(day):02d}日"

    @staticmethod
    def _apply_normal_style(doc: DocxDocument, font_name: str, size_pt: float, latin_font: str) -> None:
        normal = doc.styles["Normal"]
        normal.font.name = font_name
        normal.font.size = Pt(size_pt)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        normal.paragraph_format.line_spacing = 1.0
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)
        r_pr = normal.element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        for key in ("w:eastAsia", "w:cs"):
            r_fonts.set(qn(key), font_name)
        for key in ("w:ascii", "w:hAnsi"):
            r_fonts.set(qn(key), latin_font or font_name)

    @staticmethod
    def _normalize_all_line_spacing(doc: DocxDocument) -> None:
        paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        for paragraph in paragraphs:
            fmt = paragraph.paragraph_format
            p_pr = paragraph._p.get_or_add_pPr()
            spacing = p_pr.find(qn("w:spacing"))
            is_fixed = spacing is not None and spacing.get(qn("w:lineRule")) == "exact"
            if not is_fixed:
                fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
                fmt.line_spacing = 1.0
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)

    @staticmethod
    def _roles_from_report(report: AuditReport) -> dict[int, str]:
        roles: dict[int, str] = {}
        for element in report.elements:
            roles[element.block_index] = element.name
        return roles

    @staticmethod
    def _first_index_for(role_by_index: dict[int, str], role: str) -> int | None:
        for index, value in sorted(role_by_index.items()):
            if value == role:
                return index
        return None

    @staticmethod
    def _style_for_role(role: str) -> str:
        mapping = {
            "internal_notice": "internal_notice",
            "red_head": "red_head",
            "document_number": "document_number",
            "signer": "signer",
            "title": "title",
            "title_date": "title_date",
            "main_send": "main_send",
            "body_start": "body",
            "attachment_note": "attachment",
            "attachment_item": "attachment_item",
            "signatory": "signatory",
            "date": "date",
            "copy_to": "copy_to",
            "print_org_date": "copy_to",
            "simple_imprint": "copy_to",
            "regulation_code": "regulation_code",
            "regulation_title": "regulation_title",
            "regulation_chapter": "regulation_chapter",
            "regulation_article": "regulation_article",
            "letter_contact": "letter_contact",
            "meeting_number": "meeting_number",
            "meeting_issue_line": "meeting_issue_line",
            "meeting_attendees": "meeting_attendees",
            "distribution": "meeting_distribution",
        }
        return mapping.get(role, role if role in {"h1", "h2", "h3", "h4"} else "body")

    @staticmethod
    def _derived_role(text: str, profile: str) -> str | None:
        return heading_role(text)

    @staticmethod
    def _normalize_paragraph_text(paragraph: Paragraph) -> None:
        normalized = normalize_official_text(paragraph.text)
        if normalized != paragraph.text:
            paragraph.text = normalized

    @staticmethod
    def _force_ascii_runs_to_latin_font(paragraph: Paragraph, style) -> None:
        text = paragraph.text
        if not re.search(r"[A-Za-z0-9]", text):
            return
        parts = re.findall(r"[A-Za-z0-9]+|[^A-Za-z0-9]+", text)
        if len(parts) <= 1 and re.fullmatch(r"[A-Za-z0-9]+", text):
            for run in paragraph.runs:
                apply_font(run, style.latin_font or "Times New Roman", style.size_pt, style.color, latin_font=style.latin_font)
                if style.bold is not None:
                    run.bold = style.bold
            return
        paragraph.clear()
        for part in parts:
            run = paragraph.add_run(part)
            if re.fullmatch(r"[A-Za-z0-9]+", part):
                apply_font(run, style.latin_font or "Times New Roman", style.size_pt, style.color, latin_font=style.latin_font)
            else:
                apply_font(run, style.font, style.size_pt, style.color, latin_font=style.latin_font)
            if style.bold is not None:
                run.bold = style.bold

    @staticmethod
    def _ensure_heading_period(paragraph: Paragraph) -> None:
        text = paragraph.text.rstrip()
        if not text or text.endswith(("。", "！", "？", "；")):
            return
        for run in reversed(paragraph.runs):
            if run.text.strip():
                run.text = run.text.rstrip() + "。"
                return
        paragraph.add_run("。")

    def _apply_red_head_layout(
        self,
        role_by_index: dict[int, str],
        paragraph_by_index: dict[int, Paragraph],
    ) -> None:
        internal = self._paragraph_for_role("internal_notice", role_by_index, paragraph_by_index)
        red_head = self._paragraph_for_role("red_head", role_by_index, paragraph_by_index)
        document_number = self._paragraph_for_role("document_number", role_by_index, paragraph_by_index)

        if internal is not None and red_head is not None:
            self._set_blank_after(internal, 1)
        if red_head is not None and document_number is not None:
            self._set_blank_after(red_head, 2)
        if document_number is not None and self.config.format.draw_red_separator:
            self._clear_paragraph_border(document_number)
            self._insert_red_separator(document_number)

    def _apply_letter_head_layout(
        self,
        role_by_index: dict[int, str],
        paragraph_by_index: dict[int, Paragraph],
    ) -> None:
        """Apply letter (函) specific layout: right-align doc number."""
        document_number = self._paragraph_for_role("document_number", role_by_index, paragraph_by_index)
        if document_number is not None:
            document_number.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            letter_style = self.config.styles["letter_document_number"]
            apply_style_to_paragraph(document_number, letter_style)
            self._set_blank_after(document_number, 2)
        internal_notice = self._paragraph_for_role("internal_notice", role_by_index, paragraph_by_index)
        date = self._paragraph_for_role("date", role_by_index, paragraph_by_index)
        if internal_notice is not None and date is not None:
            self._set_blank_before(internal_notice, 2)

    def _apply_meeting_minutes_layout(
        self,
        role_by_index: dict[int, str],
        paragraph_by_index: dict[int, Paragraph],
    ) -> None:
        internal = self._paragraph_for_role("internal_notice", role_by_index, paragraph_by_index)
        red_head = self._paragraph_for_role("red_head", role_by_index, paragraph_by_index)
        number = self._paragraph_for_role("meeting_number", role_by_index, paragraph_by_index)
        issue_line = self._paragraph_for_role("meeting_issue_line", role_by_index, paragraph_by_index)
        if internal is not None:
            blank = self._set_blank_after_and_get(internal, 1)[0]
            self._set_blank_spacing(blank, 38)
        if red_head is not None:
            self._set_blank_after(red_head, 1)
        if number is not None:
            self._set_blank_after(number, 1)
        if issue_line is not None and self.config.format.draw_red_separator:
            self._clear_paragraph_border(issue_line)
            self._insert_red_separator(issue_line)

    @classmethod
    def _set_blank_after_and_get(cls, paragraph: Paragraph, count: int) -> list[Paragraph]:
        cls._set_blank_after(paragraph, count)
        result: list[Paragraph] = []
        current = paragraph._p.getnext()
        while current is not None and len(result) < count and current.tag == qn("w:p"):
            result.append(Paragraph(current, paragraph._parent))
            current = current.getnext()
        return result

    @staticmethod
    def _set_blank_spacing(paragraph: Paragraph, line_pt: float) -> None:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(line_pt)

    def _apply_regulation_layout(
        self,
        role_by_index: dict[int, str],
        paragraph_by_index: dict[int, Paragraph],
    ) -> None:
        code = self._paragraph_for_role("regulation_code", role_by_index, paragraph_by_index)
        title = self._paragraph_for_role("regulation_title", role_by_index, paragraph_by_index)
        chapters = [
            paragraph_by_index[index]
            for index, role in sorted(role_by_index.items())
            if role == "regulation_chapter" and index in paragraph_by_index
        ]
        if code is None:
            return
        self._ensure_explicit_page_break_before(code)
        self._set_blank_after(code, 1)
        if title is not None:
            self._set_blank_after(title, 1)
        for chapter in chapters:
            self._set_blank_before(chapter, 1)
            self._set_blank_after(chapter, 1)

    @staticmethod
    def _paragraph_for_role(
        role: str,
        role_by_index: dict[int, str],
        paragraph_by_index: dict[int, Paragraph],
    ) -> Paragraph | None:
        for index, value in sorted(role_by_index.items()):
            if value == role:
                return paragraph_by_index.get(index)
        return None

    def _insert_red_separator(self, document_number: Paragraph) -> None:
        current = document_number._p.getnext()
        while current is not None and current.tag == qn("w:p") and self._is_blank_paragraph_element(current):
            next_element = current.getnext()
            current.getparent().remove(current)
            current = next_element

        line_element = OxmlElement("w:p")
        document_number._p.addnext(line_element)
        line_paragraph = Paragraph(line_element, document_number._parent)
        line_paragraph.paragraph_format.space_before = Pt(0)
        line_paragraph.paragraph_format.space_after = Pt(0)
        line_paragraph.paragraph_format.line_spacing = Pt(1)
        add_floating_line(
            line_paragraph,
            width_cm=15.6,
            weight_pt=1.4,
            color=self.config.styles["red_head"].color or "FF0000",
            shape_id="OfficeToolRedSeparator",
        )
        self._set_blank_after(line_paragraph, 2)

    @staticmethod
    def _clear_paragraph_border(paragraph: Paragraph) -> None:
        p_pr = paragraph._p.pPr
        border = p_pr.find(qn("w:pBdr")) if p_pr is not None else None
        if border is not None:
            p_pr.remove(border)

    def _apply_imprint_layout(
        self,
        doc: DocxDocument,
        role_by_index: dict[int, str],
        paragraph_by_index: dict[int, Paragraph],
    ) -> None:
        copy_to = self._paragraph_for_role("copy_to", role_by_index, paragraph_by_index)
        print_line = self._paragraph_for_role("print_org_date", role_by_index, paragraph_by_index)
        if copy_to is None or print_line is None:
            return

        self._remove_generated_imprint_lines(doc)
        for paragraph in (copy_to, print_line):
            self._clear_paragraph_border(paragraph)
            clear_paragraph_frame(paragraph)

        top_line = self._new_line_paragraph_before(copy_to, "OfficeToolImprintTopLine", 1.0)
        middle_line = self._new_line_paragraph_after(copy_to, "OfficeToolImprintMiddleLine", 0.6)
        bottom_line = self._new_line_paragraph_after(print_line, "OfficeToolImprintBottomLine", 1.0)

        positions = [top_line, copy_to, middle_line, print_line, bottom_line]
        for paragraph in positions:
            clear_paragraph_frame(paragraph)
            self._clear_pagination_flags(paragraph)
        self._insert_imprint_spacers(doc, top_line)

    def _apply_distribution_layout(
        self,
        doc: DocxDocument,
        role_by_index: dict[int, str],
        paragraph_by_index: dict[int, Paragraph],
    ) -> None:
        distribution = self._paragraph_for_role("distribution", role_by_index, paragraph_by_index)
        if distribution is None:
            return
        self._remove_generated_distribution(doc)
        self._clear_paragraph_border(distribution)
        clear_paragraph_frame(distribution)
        apply_style_to_paragraph(distribution, self.config.styles["meeting_distribution"], preserve_bold_italic=False)
        top_line = self._new_line_paragraph_before(distribution, "OfficeToolDistributionTopLine", 1.0)
        bottom_line = self._new_line_paragraph_after(distribution, "OfficeToolDistributionBottomLine", 1.0)
        for paragraph in (top_line, distribution, bottom_line):
            self._clear_pagination_flags(paragraph)
        self._insert_distribution_spacers(doc, top_line)

    def _apply_simple_imprint_layout(
        self,
        doc: DocxDocument,
        role_by_index: dict[int, str],
        paragraph_by_index: dict[int, Paragraph],
    ) -> None:
        imprint = self._paragraph_for_role("simple_imprint", role_by_index, paragraph_by_index)
        if imprint is None:
            return
        self._remove_generated_simple_imprint(doc)
        self._clear_paragraph_border(imprint)
        clear_paragraph_frame(imprint)
        apply_style_to_paragraph(imprint, self.config.styles["copy_to"], preserve_bold_italic=False)
        top_line = self._new_line_paragraph_before(imprint, "OfficeToolSimpleImprintTopLine", 1.0)
        bottom_line = self._new_line_paragraph_after(imprint, "OfficeToolSimpleImprintBottomLine", 1.0)
        for paragraph in (top_line, imprint, bottom_line):
            self._clear_pagination_flags(paragraph)
        self._insert_distribution_spacers(doc, top_line)

    @staticmethod
    def _remove_generated_simple_imprint(doc: DocxDocument) -> None:
        body = doc._element.body
        tag = "{urn:schemas-microsoft-com:vml}line"
        ids = {"OfficeToolSimpleImprintTopLine", "OfficeToolSimpleImprintBottomLine"}
        for element in list(body):
            if element.tag != qn("w:p"):
                continue
            p_pr = element.find(qn("w:pPr"))
            p_style = p_pr.find(qn("w:pStyle")) if p_pr is not None else None
            if p_style is not None and p_style.get(qn("w:val")) == "OfficeToolDistributionSpacer":
                body.remove(element)
                continue
            if any(line.get("id") in ids for line in element.iter(tag)):
                body.remove(element)

    def _insert_distribution_spacers(self, doc: DocxDocument, top_line: Paragraph) -> None:
        page = 1
        line = 0
        for element in doc._element.body.iterchildren():
            if element is top_line._p:
                break
            if element.tag == qn("w:tbl"):
                page, line = self._advance_grid_position(
                    page,
                    line,
                    self._estimated_table_lines(Table(element, doc._body)),
                )
                continue
            if element.tag != qn("w:p"):
                continue
            paragraph = Paragraph(element, doc._body)
            if paragraph.paragraph_format.page_break_before and line:
                page += 1
                line = 0
            page_breaks = sum(
                1 for br in paragraph._p.iter(qn("w:br")) if br.get(qn("w:type")) == "page"
            )
            if page_breaks:
                page += page_breaks
                line = 0
                if not paragraph.text.strip():
                    continue
            page, line = self._advance_grid_position(
                page,
                line,
                self._estimated_grid_lines(paragraph),
            )

        target_page = page if page % 2 == 0 else page + 1
        target_before = (target_page - 1) * self.config.page.lines_per_page + self.config.page.lines_per_page - 4
        current = (page - 1) * self.config.page.lines_per_page + line
        while target_before < current:
            target_page += 2
            target_before = (target_page - 1) * self.config.page.lines_per_page + self.config.page.lines_per_page - 4
        for _ in range(max(0, target_before - current)):
            element = OxmlElement("w:p")
            p_pr = OxmlElement("w:pPr")
            style = OxmlElement("w:pStyle")
            style.set(qn("w:val"), "OfficeToolDistributionSpacer")
            p_pr.append(style)
            element.append(p_pr)
            top_line._p.addprevious(element)

    @staticmethod
    def _remove_generated_distribution(doc: DocxDocument) -> None:
        body = doc._element.body
        tag = "{urn:schemas-microsoft-com:vml}line"
        ids = {"OfficeToolDistributionTopLine", "OfficeToolDistributionBottomLine"}
        for element in list(body):
            if element.tag != qn("w:p"):
                continue
            p_pr = element.find(qn("w:pPr"))
            p_style = p_pr.find(qn("w:pStyle")) if p_pr is not None else None
            if p_style is not None and p_style.get(qn("w:val")) == "OfficeToolDistributionSpacer":
                body.remove(element)
                continue
            if any(line.get("id") in ids for line in element.iter(tag)):
                body.remove(element)

    def _insert_imprint_spacers(self, doc: DocxDocument, top_line: Paragraph) -> None:
        page = 1
        used_lines = 0
        for element in doc._element.body.iterchildren():
            if element is top_line._p:
                break
            if element.tag == qn("w:tbl"):
                page, used_lines = self._advance_grid_position(
                    page,
                    used_lines,
                    self._estimated_table_lines(Table(element, doc._body)),
                )
                continue
            if element.tag != qn("w:p"):
                continue
            paragraph = Paragraph(element, doc._body)
            if paragraph.paragraph_format.page_break_before and used_lines:
                page += 1
                used_lines = 0
            page_breaks = sum(
                1 for br in paragraph._p.iter(qn("w:br")) if br.get(qn("w:type")) == "page"
            )
            if page_breaks:
                page += page_breaks
                used_lines = 0
                if not paragraph.text.strip():
                    continue
            page, used_lines = self._advance_grid_position(
                page,
                used_lines,
                self._estimated_grid_lines(paragraph),
            )

        target_page = page if page % 2 == 0 else page + 1
        target_before = (target_page - 1) * self.config.page.lines_per_page + self.config.page.lines_per_page - 2
        current = (page - 1) * self.config.page.lines_per_page + used_lines
        while target_before < current:
            target_page += 2
            target_before = (target_page - 1) * self.config.page.lines_per_page + self.config.page.lines_per_page - 2
        for _ in range(max(0, target_before - current)):
            element = OxmlElement("w:p")
            p_pr = OxmlElement("w:pPr")
            style = OxmlElement("w:pStyle")
            style.set(qn("w:val"), "OfficeToolImprintSpacer")
            p_pr.append(style)
            element.append(p_pr)
            top_line._p.addprevious(element)

    def _advance_grid_position(self, page: int, line: int, added_lines: int) -> tuple[int, int]:
        line += max(0, added_lines)
        while line > self.config.page.lines_per_page:
            page += 1
            line -= self.config.page.lines_per_page
        return page, line

    def _estimated_table_lines(self, table: Table) -> int:
        total = 0
        for row in table.rows:
            cell_lines = []
            seen_cells: set[int] = set()
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                cell_lines.append(sum(self._estimated_grid_lines(p) for p in cell.paragraphs))
            total += max([1, *cell_lines])
        return max(1, total)

    def _estimated_grid_lines(self, paragraph: Paragraph) -> int:
        if not paragraph.text:
            shape_ids = {
                line.get("id")
                for line in paragraph._p.iter("{urn:schemas-microsoft-com:vml}line")
            }
            if shape_ids & {
                "OfficeToolRedSeparator",
                "OfficeToolImprintTopLine",
                "OfficeToolImprintMiddleLine",
                "OfficeToolImprintBottomLine",
                "OfficeToolDistributionTopLine",
                "OfficeToolDistributionBottomLine",
                "OfficeToolSimpleImprintTopLine",
                "OfficeToolSimpleImprintBottomLine",
            }:
                return 0
            p_pr = paragraph._p.pPr
            spacing = p_pr.find(qn("w:spacing")) if p_pr is not None else None
            if spacing is not None and spacing.get(qn("w:lineRule")) == "exact":
                line_twips = int(spacing.get(qn("w:line"), "0") or 0)
                grid_twips = int(round(self.config.page.grid_line_pitch_pt * 20))
                return max(1, (line_twips + grid_twips - 1) // grid_twips)
            return 1

        explicit_lines = paragraph.text.count("\n") + 1
        capacity = max(
            1,
            self.config.page.chars_per_line - 5
            if RE_MEETING_ATTENDEES.match(paragraph.text)
            else self.config.page.chars_per_line,
        )
        visual_lines = max(1, int((self._text_width_chars(paragraph.text.replace("\n", "")) + capacity - 1) // capacity))
        lines = max(explicit_lines, visual_lines)
        p_pr = paragraph._p.pPr
        spacing = p_pr.find(qn("w:spacing")) if p_pr is not None else None
        if spacing is not None and spacing.get(qn("w:lineRule")) == "exact":
            line_twips = int(spacing.get(qn("w:line"), "0") or 0)
            grid_twips = int(round(self.config.page.grid_line_pitch_pt * 20))
            lines *= max(1, (line_twips + grid_twips - 1) // grid_twips)
        return lines

    def _new_line_paragraph_before(self, anchor: Paragraph, shape_id: str, weight_pt: float) -> Paragraph:
        element = OxmlElement("w:p")
        anchor._p.addprevious(element)
        paragraph = Paragraph(element, anchor._parent)
        self._format_imprint_line_paragraph(paragraph, shape_id, weight_pt)
        return paragraph

    def _new_line_paragraph_after(self, anchor: Paragraph, shape_id: str, weight_pt: float) -> Paragraph:
        element = OxmlElement("w:p")
        anchor._p.addnext(element)
        paragraph = Paragraph(element, anchor._parent)
        self._format_imprint_line_paragraph(paragraph, shape_id, weight_pt)
        return paragraph

    @staticmethod
    def _format_imprint_line_paragraph(paragraph: Paragraph, shape_id: str, weight_pt: float) -> None:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(1)
        add_floating_line(
            paragraph,
            width_cm=15.6,
            weight_pt=weight_pt,
            color="000000",
            shape_id=shape_id,
        )

    @staticmethod
    def _remove_generated_imprint_lines(doc: DocxDocument) -> None:
        body = doc._element.body
        tag = "{urn:schemas-microsoft-com:vml}line"
        shape_ids = {
            "OfficeToolImprintTopLine",
            "OfficeToolImprintMiddleLine",
            "OfficeToolImprintBottomLine",
        }
        for element in list(body):
            if element.tag != qn("w:p"):
                continue
            p_pr = element.find(qn("w:pPr"))
            p_style = p_pr.find(qn("w:pStyle")) if p_pr is not None else None
            if p_style is not None and p_style.get(qn("w:val")) == "OfficeToolImprintSpacer":
                body.remove(element)
                continue
            if any(line.get("id") in shape_ids for line in element.iter(tag)):
                body.remove(element)

    @staticmethod
    def _clear_pagination_flags(paragraph: Paragraph) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        for tag in ("w:keepNext", "w:keepLines", "w:pageBreakBefore"):
            element = p_pr.find(qn(tag))
            if element is not None:
                p_pr.remove(element)

    @classmethod
    def _ensure_explicit_page_break_before(cls, paragraph: Paragraph) -> None:
        cls._clear_pagination_flags(paragraph)
        previous = paragraph._p.getprevious()
        if previous is not None and previous.tag == qn("w:p"):
            if any(
                br.get(qn("w:type")) == "page"
                for br in previous.iter(qn("w:br"))
            ):
                return
        break_paragraph = OxmlElement("w:p")
        run = OxmlElement("w:r")
        page_break = OxmlElement("w:br")
        page_break.set(qn("w:type"), "page")
        run.append(page_break)
        break_paragraph.append(run)
        paragraph._p.addprevious(break_paragraph)

    def _apply_page_number(self, doc: DocxDocument, report: AuditReport) -> None:
        style = self.config.styles["page_number"]
        set_even_and_odd_headers(doc, self.config.format.page_number_odd_even)
        for section_index, section in enumerate(doc.sections):
            if report.is_letter_head and section_index == 0:
                section.different_first_page_header_footer = True
                first_footer = section.first_page_footer
                for paragraph in first_footer.paragraphs:
                    paragraph.text = ""
            footer = section.footer
            paragraph: Paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            alignment = WD_ALIGN_PARAGRAPH.RIGHT if self.config.format.page_number_odd_even else WD_ALIGN_PARAGRAPH.CENTER
            edge_space = 1 if self.config.format.page_number_odd_even else 0
            add_page_number(paragraph, style.font, style.size_pt, alignment=alignment, with_dashes=True, edge_space_chars=edge_space)
            if self.config.format.page_number_odd_even:
                even_footer = section.even_page_footer
                even_paragraph: Paragraph = even_footer.paragraphs[0] if even_footer.paragraphs else even_footer.add_paragraph()
                add_page_number(
                    even_paragraph,
                    style.font,
                    style.size_pt,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    with_dashes=True,
                    edge_space_chars=1,
                )

    def _apply_letter_header(self, doc: DocxDocument, report: AuditReport) -> None:
        """Apply the first-page letter head, compound rules, and hidden first-page number."""
        red_head_element = next((element for element in report.elements if element.name == "red_head"), None)
        if red_head_element is None:
            return
        red_head_paragraph = doc.paragraphs[red_head_element.block_index]
        red_head_text = red_head_paragraph.text.strip()
        if not red_head_text:
            return
        section = doc.sections[0]
        set_even_and_odd_headers(doc, True)
        letter_style = self.config.styles["letter_red_head"]
        font_size, character_spacing = self._fit_letter_red_head(red_head_text, letter_style.size_pt)
        add_body_text_box(
            red_head_paragraph,
            red_head_text,
            letter_style.font,
            font_size,
            letter_style.color or "FF0000",
            letter_style.line_spacing_pt or 40,
            character_spacing_pt=character_spacing,
            top_cm=3.0,
        )
        red_head_paragraph.paragraph_format.space_before = Pt(0)
        red_head_paragraph.paragraph_format.space_after = Pt(0)
        red_head_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        red_head_paragraph.paragraph_format.line_spacing = Pt(24)
        line_element = OxmlElement("w:p")
        red_head_paragraph._p.addnext(line_element)
        top_line_paragraph = Paragraph(line_element, red_head_paragraph._parent)
        self._format_letter_red_line(
            top_line_paragraph,
            "OfficeToolLetterRedTop",
            "thickThin",
        )
        add_floating_line(
            red_head_paragraph,
            width_cm=17.0,
            weight_pt=4.0,
            color="FF0000",
            shape_id="OfficeToolLetterRedBottom",
            compound="thinThick",
            vertical_offset_pt=(29.7 - 2.5) / 2.54 * 72,
            vertical_relative="page",
            clear_paragraph=False,
        )

    def _fit_letter_red_head(self, text: str, requested_size_pt: float) -> tuple[float, float]:
        units = max(1.0, self._text_width_chars(text.replace("\n", "")))
        gaps = max(1, len(text) - 1)
        available_pt = 15.5 / 2.54 * 72 - 4
        font_size = min(41.0, requested_size_pt or 41.0)
        spacing_pt = 0.0
        if units * font_size > available_pt:
            spacing_pt = (available_pt - units * font_size) / gaps
            if spacing_pt < -0.8:
                spacing_pt = -0.8
                font_size = max(22.0, int(((available_pt - spacing_pt * gaps) / units) * 2) / 2)
        return font_size, spacing_pt

    @staticmethod
    def _format_letter_red_line(
        paragraph: Paragraph,
        shape_id: str,
        compound: str,
        *,
        vertical_offset_pt: float = 0.0,
        line_height_pt: float = 1.0,
    ) -> None:
        """Format a 17 cm, 4 pt compound letter-head rule."""
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(line_height_pt)
        add_floating_line(
            paragraph,
            width_cm=17.0,
            weight_pt=4.0,
            color="FF0000",
            shape_id=shape_id,
            compound=compound,
            vertical_offset_pt=vertical_offset_pt,
        )


HALFWIDTH_TO_FULLWIDTH = str.maketrans(
    {
        ":": "：",
        ",": "，",
        ";": "；",
        "!": "！",
        "?": "？",
        "(": "（",
        ")": "）",
    }
)


def normalize_official_text(text: str) -> str:
    if not text:
        return text
    text = re.sub(r"(?m)^(\s*\d{1,2})\s+\.", r"\1.", text)
    text = re.sub(r"(?m)^(\s*\d{1,2})\.(?=\s*)", r"\1．", text)
    text = text.translate(HALFWIDTH_TO_FULLWIDTH)
    text = re.sub(r"(?m)^(\s*\d{1,2}[．、])\s+", r"\1", text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])", "", text)
    text = re.sub(r"\s+([：，；！？。、）])", r"\1", text)
    text = re.sub(r"([（])\s+", r"\1", text)
    text = re.sub(r"([：，；！？。、])\s+", r"\1", text)
    text = re.sub(r" {2,}", " ", text)
    return text
